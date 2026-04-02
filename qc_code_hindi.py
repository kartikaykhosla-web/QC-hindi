# -*- coding: utf-8 -*-
"""
QC Code SSG — Hindi (STRUCTURAL PARITY VERSION)
Vertex AI Gemini 2.5 Flash

Goal:
- Same logical scaffolding as English QC
- Same safety rails
- Same determinism philosophy
- Hindi-compatible engines
"""

# =================================================
# GLOBAL CACHE
# =================================================
FACT_CACHE = {}

# =================================================
# CORE IMPORTS
# =================================================
import re
import os
import json
import base64
import requests
import hashlib
import tempfile
import io
import html
import unicodedata
import sqlite3
import uuid
from datetime import datetime, timezone, timedelta
import streamlit as st
IMPORT_ONLY = os.environ.get("QC_HINDI_IMPORT_ONLY") == "1"
try:
    import extra_streamlit_components as stx
except Exception:
    if IMPORT_ONLY:
        class _DummyCookieManager:
            def get(self, *args, **kwargs):
                return ""

            def set(self, *args, **kwargs):
                return None

            def delete(self, *args, **kwargs):
                return None

        class _DummySTX:
            CookieManager = _DummyCookieManager

        stx = _DummySTX()
    else:
        raise
from bs4 import BeautifulSoup
from google.oauth2 import service_account
from google.auth.transport.requests import Request as GoogleAuthRequest
from difflib import SequenceMatcher
from urllib.parse import urlparse, quote

# =================================================
# GEN AI CLIENT
# =================================================
from google import genai
from google.genai import types as genai_types

# =================================================
# STREAMLIT CONFIG
# =================================================
if not IMPORT_ONLY:
    st.set_page_config(page_title="Hindi Article QC Tool (Gemini)", layout="wide")

def _secret(name: str, default=""):
    try:
        return st.secrets[name]
    except Exception:
        return default

ALLOWED_EMAIL_DOMAIN = str(_secret("ALLOWED_EMAIL_DOMAIN", "jagrannewmedia.com")).strip().lower()
APP_ACCESS_SUPPORT_TEXT = str(
    _secret(
        "APP_ACCESS_SUPPORT_TEXT",
        "Enter your Jagran New Media email address to continue.",
    )
).strip()
ADMIN_EMAIL = "kartikay.khosla@jagrannewmedia.com"
HISTORY_DB_PATH = os.path.join(os.path.dirname(__file__), ".app_history.sqlite3")
HISTORY_SPREADSHEET_ID = str(_secret("HISTORY_SPREADSHEET_ID", "")).strip()
SESSION_QUERY_KEY = "_jnm_session"
SESSION_COOKIE_KEY = "_jnm_session"
SESSION_TTL_HOURS = 24
SESSION_REFRESH_WINDOW_MINUTES = 15

HISTORY_HEADERS = {
    "login_events": ["ts_utc", "app", "email"],
    "analysis_runs": [
        "run_id",
        "ts_utc",
        "app",
        "email",
        "source_type",
        "source_identity",
        "source_label",
        "analysis_key",
        "iteration",
        "spelling_count",
        "grammar_count",
        "editorial_count",
        "fact_count",
        "total_count",
    ],
    "access_sessions": [
        "ts_utc",
        "app",
        "token_hash",
        "email",
        "event_type",
        "expires_ts_utc",
        "last_seen_ts_utc",
    ],
}

def _email_allowed(email: str) -> bool:
    return (email or "").strip().lower().endswith(f"@{ALLOWED_EMAIL_DOMAIN}")

def _normalise_username(username: str) -> str:
    value = (username or "").strip().lower().replace(" ", "")
    if "@" in value:
        return ""
    return value

def _build_email_from_username(username: str) -> str:
    username = _normalise_username(username)
    if not username:
        return ""
    return f"{username}@{ALLOWED_EMAIL_DOMAIN}"

def _email_access_granted() -> bool:
    email = st.session_state.get("_email_access_email", "")
    return bool(st.session_state.get("_email_access_granted")) and _email_allowed(email)

def _clear_email_access():
    st.session_state.pop("_email_access_granted", None)
    st.session_state.pop("_email_access_email", None)
    _clear_pending_analysis_state()

def _current_access_email() -> str:
    return (st.session_state.get("_email_access_email") or "").strip().lower()

def _is_admin_user() -> bool:
    return _current_access_email() == ADMIN_EMAIL

def _history_conn():
    conn = sqlite3.connect(HISTORY_DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def _history_uses_sheets() -> bool:
    return bool(HISTORY_SPREADSHEET_ID)

def _sheet_tab_title(app_name: str, kind: str) -> str:
    safe_app = re.sub(r"[^A-Za-z0-9_\\-]", "_", app_name or "app").strip("_") or "app"
    safe_kind = re.sub(r"[^A-Za-z0-9_\\-]", "_", kind or "history").strip("_") or "history"
    return f"{safe_app}_{safe_kind}"[:95]

def _history_headers(kind: str):
    return HISTORY_HEADERS.get(kind, [])

def _parse_utc_iso(value: str):
    text = (value or "").strip()
    if not text:
        return None
    if text.endswith("Z"):
        text = text[:-1] + "+00:00"
    try:
        parsed = datetime.fromisoformat(text)
        if parsed.tzinfo is None:
            parsed = parsed.replace(tzinfo=timezone.utc)
        return parsed.astimezone(timezone.utc)
    except Exception:
        return None

def _safe_int(value, default=0):
    try:
        return int(value)
    except Exception:
        return default

def _utc_now():
    return datetime.now(timezone.utc)

def _hash_session_token(token: str) -> str:
    return hashlib.sha256(f"{ALLOWED_EMAIL_DOMAIN}:{token}".encode("utf-8")).hexdigest()

def _get_cookie_manager():
    manager = st.session_state.get("_cookie_manager")
    if manager is None:
        manager = stx.CookieManager()
        st.session_state["_cookie_manager"] = manager
    return manager

def _get_session_query_token() -> str:
    try:
        value = st.query_params.get(SESSION_QUERY_KEY, "")
        if isinstance(value, list):
            return (value[0] or "").strip()
        return (value or "").strip()
    except Exception:
        return ""

def _set_session_query_token(token: str):
    try:
        st.query_params[SESSION_QUERY_KEY] = token
    except Exception:
        pass

def _clear_session_query_token():
    try:
        st.query_params.pop(SESSION_QUERY_KEY, None)
    except Exception:
        pass

def _get_session_cookie_token() -> str:
    try:
        value = _get_cookie_manager().get(SESSION_COOKIE_KEY)
        return (value or "").strip()
    except Exception:
        return ""

def _set_session_cookie_token(token: str):
    try:
        _get_cookie_manager().set(
            SESSION_COOKIE_KEY,
            token,
            expires_at=datetime.now() + timedelta(hours=SESSION_TTL_HOURS),
            key=f"set-cookie-{SESSION_COOKIE_KEY}",
        )
    except Exception:
        pass

def _clear_session_cookie_token():
    try:
        _get_cookie_manager().delete(SESSION_COOKIE_KEY, key=f"delete-cookie-{SESSION_COOKIE_KEY}")
    except Exception:
        pass

def _sqlite_rows(query: str, params=()):
    try:
        ensure_history_db()
        with _history_conn() as conn:
            return [dict(row) for row in conn.execute(query, params).fetchall()]
    except Exception:
        return []

def _load_service_account_info():
    if "GCP_SERVICE_ACCOUNT_JSON_B64" not in st.secrets:
        st.error("❌ GCP_SERVICE_ACCOUNT_JSON_B64 missing")
        st.stop()

    decoded = base64.b64decode(
        st.secrets["GCP_SERVICE_ACCOUNT_JSON_B64"]
    ).decode("utf-8")
    return json.loads(decoded)

def _get_scoped_service_account_credentials(scopes):
    creds_dict = _load_service_account_info()
    with open(CRED_PATH, "w") as f:
        json.dump(creds_dict, f)
    os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = CRED_PATH
    creds = service_account.Credentials.from_service_account_info(
        creds_dict,
        scopes=list(scopes),
    )
    project_id = PROJECT_ID or str(creds_dict.get("project_id", "")).strip()
    if not project_id:
        st.error("❌ Could not determine Vertex project ID from secrets or service account JSON")
        st.stop()
    return creds, project_id, creds_dict

def _sheets_api_request(method: str, path: str = "", params=None, json_body=None):
    creds, _, _ = _get_scoped_service_account_credentials([CLOUD_PLATFORM_SCOPE, SHEETS_SCOPE])
    creds.refresh(GoogleAuthRequest())
    headers = {"Authorization": f"Bearer {creds.token}"}
    if json_body is not None:
        headers["Content-Type"] = "application/json"
    url = f"https://sheets.googleapis.com/v4/spreadsheets/{HISTORY_SPREADSHEET_ID}{path}"
    response = requests.request(
        method,
        url,
        headers=headers,
        params=params,
        json=json_body,
        timeout=30,
    )
    if response.status_code >= 400:
        raise RuntimeError(f"Google Sheets API {response.status_code}: {response.text}")
    if not response.text:
        return {}
    return response.json()

def _ensure_history_sheet(app_name: str, kind: str) -> str:
    if not _history_uses_sheets():
        return ""

    tab_title = _sheet_tab_title(app_name, kind)
    metadata = _sheets_api_request("GET", params={"fields": "sheets.properties.title"})
    titles = {
        ((sheet.get("properties") or {}).get("title") or "").strip()
        for sheet in metadata.get("sheets", [])
    }
    if tab_title not in titles:
        _sheets_api_request(
            "POST",
            ":batchUpdate",
            json_body={"requests": [{"addSheet": {"properties": {"title": tab_title}}}]},
        )

    headers = _history_headers(kind)
    encoded_range = quote(f"{tab_title}!1:1", safe="!:$")
    current = _sheets_api_request("GET", f"/values/{encoded_range}")
    current_values = current.get("values", [])
    if not current_values:
        encoded_write = quote(f"{tab_title}!A1", safe="!:$")
        _sheets_api_request(
            "PUT",
            f"/values/{encoded_write}",
            params={"valueInputOption": "RAW"},
            json_body={"range": f"{tab_title}!A1", "majorDimension": "ROWS", "values": [headers]},
        )
    return tab_title

def _sheet_read_rows(app_name: str, kind: str):
    if not _history_uses_sheets():
        return []
    try:
        tab_title = _ensure_history_sheet(app_name, kind)
        encoded_range = quote(f"{tab_title}!A:Z", safe="!:$")
        payload = _sheets_api_request("GET", f"/values/{encoded_range}")
        values = payload.get("values", [])
        headers = _history_headers(kind)
        if not values:
            return []
        start_index = 1 if values[0] == headers else 0
        rows = []
        for raw_row in values[start_index:]:
            if not any((cell or "").strip() for cell in raw_row):
                continue
            padded = list(raw_row) + [""] * max(0, len(headers) - len(raw_row))
            rows.append({header: padded[idx] if idx < len(padded) else "" for idx, header in enumerate(headers)})
        return rows
    except Exception:
        return []

def _sheet_append_row(app_name: str, kind: str, row_dict: dict) -> bool:
    if not _history_uses_sheets():
        return False
    try:
        tab_title = _ensure_history_sheet(app_name, kind)
        headers = _history_headers(kind)
        encoded_range = quote(f"{tab_title}!A:Z", safe="!:$")
        _sheets_api_request(
            "POST",
            f"/values/{encoded_range}:append",
            params={"valueInputOption": "RAW", "insertDataOption": "INSERT_ROWS"},
            json_body={"values": [[str(row_dict.get(header, "")) for header in headers]]},
        )
        return True
    except Exception:
        return False

def _sheet_login_rows(app_name: str):
    rows = _sheet_read_rows(app_name, "login_events")
    return sorted(
        [
            {
                "ts_utc": (row.get("ts_utc") or "").strip(),
                "email": (row.get("email") or "").strip().lower(),
            }
            for row in rows
            if (row.get("email") or "").strip()
        ],
        key=lambda row: row.get("ts_utc", ""),
        reverse=True,
    )

def _sheet_analysis_rows(app_name: str):
    rows = []
    for row in _sheet_read_rows(app_name, "analysis_runs"):
        ts_value = (row.get("ts_utc") or "").strip()
        email = (row.get("email") or "").strip().lower()
        if not ts_value or not email:
            continue
        rows.append(
            {
                "run_id": (row.get("run_id") or "").strip(),
                "ts_utc": ts_value,
                "app": (row.get("app") or "").strip(),
                "email": email,
                "source_type": (row.get("source_type") or "").strip(),
                "source_identity": (row.get("source_identity") or "").strip(),
                "source_label": (row.get("source_label") or "").strip(),
                "analysis_key": (row.get("analysis_key") or "").strip(),
                "iteration": _safe_int(row.get("iteration")),
                "spelling_count": _safe_int(row.get("spelling_count")),
                "grammar_count": _safe_int(row.get("grammar_count")),
                "editorial_count": _safe_int(row.get("editorial_count")),
                "fact_count": _safe_int(row.get("fact_count")),
                "total_count": _safe_int(row.get("total_count")),
            }
        )
    return sorted(rows, key=lambda row: row.get("ts_utc", ""), reverse=True)

def _sheet_latest_session_row(app_name: str, token_hash: str):
    rows = [
        row
        for row in _sheet_read_rows(app_name, "access_sessions")
        if (row.get("token_hash") or "").strip() == token_hash
    ]
    if not rows:
        return None
    rows.sort(key=lambda row: (row.get("ts_utc") or "", row.get("last_seen_ts_utc") or ""), reverse=True)
    latest = rows[0]
    return {
        "ts_utc": (latest.get("ts_utc") or "").strip(),
        "app": (latest.get("app") or "").strip(),
        "token_hash": (latest.get("token_hash") or "").strip(),
        "email": (latest.get("email") or "").strip().lower(),
        "event_type": (latest.get("event_type") or "").strip().lower(),
        "expires_ts_utc": (latest.get("expires_ts_utc") or "").strip(),
        "last_seen_ts_utc": (latest.get("last_seen_ts_utc") or "").strip(),
    }

def _append_session_event(app_name: str, token_hash: str, email: str, event_type: str, expires_iso: str, last_seen_iso: str) -> bool:
    return _sheet_append_row(
        app_name,
        "access_sessions",
        {
            "ts_utc": _utc_now().isoformat(),
            "app": app_name,
            "token_hash": token_hash,
            "email": (email or "").strip().lower(),
            "event_type": event_type,
            "expires_ts_utc": expires_iso,
            "last_seen_ts_utc": last_seen_iso,
        },
    )

def ensure_history_db():
    try:
        with _history_conn() as conn:
            conn.execute(
                """
                CREATE TABLE IF NOT EXISTS login_events (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    ts_utc TEXT NOT NULL,
                    app TEXT NOT NULL,
                    email TEXT NOT NULL
                )
                """
            )
            conn.execute(
                """
                CREATE TABLE IF NOT EXISTS analysis_runs (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    run_id TEXT NOT NULL UNIQUE,
                    ts_utc TEXT NOT NULL,
                    app TEXT NOT NULL,
                    email TEXT NOT NULL,
                    source_type TEXT NOT NULL,
                    source_identity TEXT NOT NULL,
                    source_label TEXT NOT NULL,
                    analysis_key TEXT,
                    iteration INTEGER NOT NULL,
                    spelling_count INTEGER NOT NULL DEFAULT 0,
                    grammar_count INTEGER NOT NULL DEFAULT 0,
                    editorial_count INTEGER NOT NULL DEFAULT 0,
                    fact_count INTEGER NOT NULL DEFAULT 0,
                    total_count INTEGER NOT NULL DEFAULT 0
                )
                """
            )
            conn.execute(
                """
                CREATE TABLE IF NOT EXISTS access_sessions (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    token_hash TEXT NOT NULL UNIQUE,
                    app TEXT NOT NULL,
                    email TEXT NOT NULL,
                    created_ts_utc TEXT NOT NULL,
                    expires_ts_utc TEXT NOT NULL,
                    revoked INTEGER NOT NULL DEFAULT 0,
                    last_seen_ts_utc TEXT
                )
                """
            )
            conn.execute(
                "CREATE INDEX IF NOT EXISTS idx_login_events_app_email_ts ON login_events(app, email, ts_utc)"
            )
            conn.execute(
                "CREATE INDEX IF NOT EXISTS idx_analysis_runs_app_source_ts ON analysis_runs(app, source_identity, ts_utc)"
            )
            conn.execute(
                "CREATE INDEX IF NOT EXISTS idx_access_sessions_app_email_expiry ON access_sessions(app, email, expires_ts_utc)"
            )
    except Exception:
        pass

def _clear_pending_analysis_state():
    for key in (
        "_pending_run_id",
        "_pending_source_type",
        "_pending_source_identity",
        "_pending_source_label",
        "_pending_analysis_key",
    ):
        st.session_state.pop(key, None)

def _create_persisted_session(app_name: str, email: str):
    try:
        token = uuid.uuid4().hex + uuid.uuid4().hex
        now_iso = _utc_now().isoformat()
        expires_iso = (_utc_now() + timedelta(hours=SESSION_TTL_HOURS)).isoformat()
        token_hash = _hash_session_token(token)
        stored = False
        if _history_uses_sheets():
            stored = _append_session_event(app_name, token_hash, email, "create", expires_iso, now_iso)
        if not stored:
            ensure_history_db()
            with _history_conn() as conn:
                conn.execute(
                    """
                    INSERT INTO access_sessions (
                        token_hash, app, email, created_ts_utc, expires_ts_utc, revoked, last_seen_ts_utc
                    ) VALUES (?, ?, ?, ?, ?, 0, ?)
                    """,
                    (token_hash, app_name, (email or "").strip().lower(), now_iso, expires_iso, now_iso),
                )
        _set_session_query_token(token)
        _set_session_cookie_token(token)
    except Exception:
        pass

def _revoke_persisted_session(app_name: str):
    token = _get_session_query_token() or _get_session_cookie_token()
    if token:
        try:
            token_hash = _hash_session_token(token)
            stored = False
            if _history_uses_sheets():
                stored = _append_session_event(
                    app_name,
                    token_hash,
                    _current_access_email(),
                    "revoke",
                    _utc_now().isoformat(),
                    _utc_now().isoformat(),
                )
            if not stored:
                ensure_history_db()
                with _history_conn() as conn:
                    conn.execute(
                        "UPDATE access_sessions SET revoked = 1 WHERE app = ? AND token_hash = ?",
                        (app_name, token_hash),
                    )
        except Exception:
            pass
    _clear_session_query_token()
    _clear_session_cookie_token()

def _restore_persisted_session(app_name: str) -> bool:
    if _email_access_granted():
        return True

    token = _get_session_query_token() or _get_session_cookie_token()
    if not token:
        return False

    try:
        now_iso = _utc_now().isoformat()
        token_hash = _hash_session_token(token)
        row = None
        if _history_uses_sheets():
            row = _sheet_latest_session_row(app_name, token_hash)
            if not row:
                _clear_session_query_token()
                _clear_session_cookie_token()
                return False
            if row.get("event_type") == "revoke":
                _clear_session_query_token()
                _clear_session_cookie_token()
                return False
            expires_at = _parse_utc_iso(row.get("expires_ts_utc", ""))
            if not expires_at or expires_at <= _utc_now():
                _clear_session_query_token()
                _clear_session_cookie_token()
                return False
            refreshed_expiry = (_utc_now() + timedelta(hours=SESSION_TTL_HOURS)).isoformat()
            last_seen = _parse_utc_iso(row.get("last_seen_ts_utc", ""))
            if (not last_seen) or ((_utc_now() - last_seen) >= timedelta(minutes=SESSION_REFRESH_WINDOW_MINUTES)):
                _append_session_event(
                    app_name,
                    token_hash,
                    row.get("email", ""),
                    "refresh",
                    refreshed_expiry,
                    now_iso,
                )
        else:
            ensure_history_db()
            with _history_conn() as conn:
                sqlite_row = conn.execute(
                    """
                    SELECT email
                    FROM access_sessions
                    WHERE app = ?
                      AND token_hash = ?
                      AND revoked = 0
                      AND expires_ts_utc > ?
                    ORDER BY id DESC
                    LIMIT 1
                    """,
                    (app_name, token_hash, now_iso),
                ).fetchone()
                if not sqlite_row:
                    _clear_session_query_token()
                    _clear_session_cookie_token()
                    return False

                refreshed_expiry = (_utc_now() + timedelta(hours=SESSION_TTL_HOURS)).isoformat()
                conn.execute(
                    """
                    UPDATE access_sessions
                    SET last_seen_ts_utc = ?, expires_ts_utc = ?
                    WHERE app = ? AND token_hash = ?
                    """,
                    (now_iso, refreshed_expiry, app_name, token_hash),
                )
            row = {"email": sqlite_row["email"]}

        email = (row.get("email") or "").strip().lower()
        if not _email_allowed(email):
            _clear_session_query_token()
            _clear_session_cookie_token()
            return False

        _set_session_query_token(token)
        _set_session_cookie_token(token)
        st.session_state["_email_access_granted"] = True
        st.session_state["_email_access_email"] = email
        return True
    except Exception:
        return False

def queue_analysis_run(source_type: str, source_identity: str, source_label: str, analysis_key: str = ""):
    st.session_state["_pending_run_id"] = uuid.uuid4().hex
    st.session_state["_pending_source_type"] = source_type
    st.session_state["_pending_source_identity"] = source_identity
    st.session_state["_pending_source_label"] = source_label
    st.session_state["_pending_analysis_key"] = analysis_key

def _record_access_event(app_name: str, email: str):
    try:
        row = {
            "ts_utc": datetime.now(timezone.utc).isoformat(),
            "app": app_name,
            "email": (email or "").strip().lower(),
        }
        stored = _sheet_append_row(app_name, "login_events", row) if _history_uses_sheets() else False
        if not stored:
            ensure_history_db()
            with _history_conn() as conn:
                conn.execute(
                    "INSERT INTO login_events (ts_utc, app, email) VALUES (?, ?, ?)",
                    (row["ts_utc"], row["app"], row["email"]),
                )
    except Exception:
        pass

def count_markdown_rows(table_md: str, header_name: str) -> int:
    count = 0
    for line in (table_md or "").splitlines():
        row = line.strip()
        if not row.startswith("|") or row.count("|") < 2:
            continue
        parts = [part.strip() for part in row.strip("|").split("|")]
        if not parts or not any(parts):
            continue
        if all(re.fullmatch(r":?-{2,}:?", part or "") for part in parts):
            continue
        if parts[0].lower() == header_name.lower():
            continue
        count += 1
    return count

def compute_qc_score(spelling_count: int, grammar_count: int, editorial_count: int, fact_count: int) -> int:
    weighted_penalty = (
        float(spelling_count) * 0.5
        + float(grammar_count) * 1.0
        + float(editorial_count) * 0.75
        + float(fact_count) * 4.0
    )
    return max(0, min(100, int(round(100 - min(100, weighted_penalty)))))

def render_qc_score_summary(spelling_count: int, grammar_count: int, editorial_count: int, fact_count: int, has_ai_error: bool):
    st.markdown("### QC Summary")
    if has_ai_error:
        st.warning("QC score is unavailable because one or more AI checks failed.")
        return
    total_count = spelling_count + grammar_count + editorial_count + fact_count
    score = compute_qc_score(spelling_count, grammar_count, editorial_count, fact_count)
    score_col, spelling_col, grammar_col, editorial_col, fact_col, total_col = st.columns(6)
    score_col.metric("QC Score", f"{score}/100")
    spelling_col.metric("Spelling", spelling_count)
    grammar_col.metric("Grammar", grammar_count)
    editorial_col.metric("Editorial", editorial_count)
    fact_col.metric("Fact", fact_count)
    total_col.metric("Total Issues", total_count)
    st.caption("QC score is a weighted indicator based on issue counts. Fact issues carry the highest penalty.")

def log_analysis_run(app_name: str, email: str, source_type: str, source_identity: str, source_label: str,
                     analysis_key: str, spelling_count: int, grammar_count: int,
                     editorial_count: int, fact_count: int):
    run_id = st.session_state.get("_pending_run_id")
    if not run_id:
        return

    try:
        ts_utc = datetime.now(timezone.utc).isoformat()
        if _history_uses_sheets():
            existing_rows = _sheet_analysis_rows(app_name)
            if any((row.get("run_id") or "").strip() == run_id for row in existing_rows):
                _clear_pending_analysis_state()
                return
            iteration = max(
                [row.get("iteration", 0) for row in existing_rows if (row.get("source_identity") or "") == source_identity] or [0]
            ) + 1
            stored = _sheet_append_row(
                app_name,
                "analysis_runs",
                {
                    "run_id": run_id,
                    "ts_utc": ts_utc,
                    "app": app_name,
                    "email": (email or "").strip().lower(),
                    "source_type": source_type,
                    "source_identity": source_identity,
                    "source_label": source_label,
                    "analysis_key": analysis_key,
                    "iteration": iteration,
                    "spelling_count": int(spelling_count),
                    "grammar_count": int(grammar_count),
                    "editorial_count": int(editorial_count),
                    "fact_count": int(fact_count),
                    "total_count": int(spelling_count + grammar_count + editorial_count + fact_count),
                },
            )
            if not stored:
                raise RuntimeError("sheets-write-failed")
        else:
            ensure_history_db()
            with _history_conn() as conn:
                exists = conn.execute(
                    "SELECT 1 FROM analysis_runs WHERE run_id = ?",
                    (run_id,),
                ).fetchone()
                if exists:
                    _clear_pending_analysis_state()
                    return

                iteration = conn.execute(
                    "SELECT COALESCE(MAX(iteration), 0) + 1 FROM analysis_runs WHERE app = ? AND source_identity = ?",
                    (app_name, source_identity),
                ).fetchone()[0]

                conn.execute(
                    """
                    INSERT INTO analysis_runs (
                        run_id, ts_utc, app, email, source_type, source_identity, source_label,
                        analysis_key, iteration, spelling_count, grammar_count, editorial_count,
                        fact_count, total_count
                    ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        run_id,
                        ts_utc,
                        app_name,
                        (email or "").strip().lower(),
                        source_type,
                        source_identity,
                        source_label,
                        analysis_key,
                        iteration,
                        int(spelling_count),
                        int(grammar_count),
                        int(editorial_count),
                        int(fact_count),
                        int(spelling_count + grammar_count + editorial_count + fact_count),
                    ),
                )
        _clear_pending_analysis_state()
    except Exception:
        pass

def _fetch_rows(query: str, params=()):
    return _sqlite_rows(query, params)

def render_admin_dashboard(app_name: str):
    st.divider()
    with st.expander("Admin History", expanded=False):
        if _history_uses_sheets():
            login_rows = _sheet_login_rows(app_name)[:200]
            all_analysis_rows = _sheet_analysis_rows(app_name)
            recent_rows = [
                {
                    "ts_utc": row["ts_utc"],
                    "email": row["email"],
                    "source_type": row["source_type"],
                    "source_label": row["source_label"],
                    "iteration": row["iteration"],
                    "spelling_count": row["spelling_count"],
                    "grammar_count": row["grammar_count"],
                    "editorial_count": row["editorial_count"],
                    "fact_count": row["fact_count"],
                    "total_count": row["total_count"],
                }
                for row in all_analysis_rows[:200]
            ]

            daily_map = {}
            for row in login_rows:
                date_utc = (row.get("ts_utc") or "")[:10]
                email = (row.get("email") or "").strip().lower()
                if not date_utc or not email:
                    continue
                entry = daily_map.setdefault(
                    (date_utc, email),
                    {
                        "date_utc": date_utc,
                        "email": email,
                        "login_count": 0,
                        "analyses_run": 0,
                        "distinct_articles": set(),
                        "spelling_issues": 0,
                        "grammar_issues": 0,
                        "editorial_issues": 0,
                        "fact_issues": 0,
                    },
                )
                entry["login_count"] += 1

            for row in all_analysis_rows:
                date_utc = (row.get("ts_utc") or "")[:10]
                email = (row.get("email") or "").strip().lower()
                if not date_utc or not email:
                    continue
                entry = daily_map.setdefault(
                    (date_utc, email),
                    {
                        "date_utc": date_utc,
                        "email": email,
                        "login_count": 0,
                        "analyses_run": 0,
                        "distinct_articles": set(),
                        "spelling_issues": 0,
                        "grammar_issues": 0,
                        "editorial_issues": 0,
                        "fact_issues": 0,
                    },
                )
                entry["analyses_run"] += 1
                entry["distinct_articles"].add((row.get("source_identity") or "").strip())
                entry["spelling_issues"] += _safe_int(row.get("spelling_count"))
                entry["grammar_issues"] += _safe_int(row.get("grammar_count"))
                entry["editorial_issues"] += _safe_int(row.get("editorial_count"))
                entry["fact_issues"] += _safe_int(row.get("fact_count"))

            daily_rows = []
            for entry in daily_map.values():
                daily_rows.append(
                    {
                        "date_utc": entry["date_utc"],
                        "email": entry["email"],
                        "login_count": entry["login_count"],
                        "analyses_run": entry["analyses_run"],
                        "distinct_articles": len([value for value in entry["distinct_articles"] if value]),
                        "spelling_issues": entry["spelling_issues"],
                        "grammar_issues": entry["grammar_issues"],
                        "editorial_issues": entry["editorial_issues"],
                        "fact_issues": entry["fact_issues"],
                    }
                )
            daily_rows.sort(key=lambda row: (row["date_utc"], row["email"]), reverse=True)
            daily_rows = daily_rows[:180]
        else:
            daily_rows = _fetch_rows(
                """
                WITH login_daily AS (
                    SELECT
                        substr(ts_utc, 1, 10) AS date_utc,
                        email,
                        COUNT(*) AS login_count
                    FROM login_events
                    WHERE app = ?
                    GROUP BY substr(ts_utc, 1, 10), email
                ),
                analysis_daily AS (
                    SELECT
                        substr(ts_utc, 1, 10) AS date_utc,
                        email,
                        COUNT(*) AS analyses_run,
                        COUNT(DISTINCT source_identity) AS distinct_articles,
                        COALESCE(SUM(spelling_count), 0) AS spelling_issues,
                        COALESCE(SUM(grammar_count), 0) AS grammar_issues,
                        COALESCE(SUM(editorial_count), 0) AS editorial_issues,
                        COALESCE(SUM(fact_count), 0) AS fact_issues
                    FROM analysis_runs
                    WHERE app = ?
                    GROUP BY substr(ts_utc, 1, 10), email
                ),
                combined AS (
                    SELECT date_utc, email FROM login_daily
                    UNION
                    SELECT date_utc, email FROM analysis_daily
                )
                SELECT
                    c.date_utc,
                    c.email,
                    COALESCE(l.login_count, 0) AS login_count,
                    COALESCE(a.analyses_run, 0) AS analyses_run,
                    COALESCE(a.distinct_articles, 0) AS distinct_articles,
                    COALESCE(a.spelling_issues, 0) AS spelling_issues,
                    COALESCE(a.grammar_issues, 0) AS grammar_issues,
                    COALESCE(a.editorial_issues, 0) AS editorial_issues,
                    COALESCE(a.fact_issues, 0) AS fact_issues
                FROM combined c
                LEFT JOIN login_daily l
                  ON l.date_utc = c.date_utc AND l.email = c.email
                LEFT JOIN analysis_daily a
                  ON a.date_utc = c.date_utc AND a.email = c.email
                ORDER BY c.date_utc DESC, c.email ASC
                LIMIT 180
                """,
                (app_name, app_name),
            )

            login_rows = _fetch_rows(
                """
                SELECT ts_utc, email
                FROM login_events
                WHERE app = ?
                ORDER BY ts_utc DESC
                LIMIT 200
                """,
                (app_name,),
            )

            recent_rows = _fetch_rows(
                """
                SELECT
                    ts_utc,
                    email,
                    source_type,
                    source_label,
                    iteration,
                    spelling_count,
                    grammar_count,
                    editorial_count,
                    fact_count,
                    total_count
                FROM analysis_runs
                WHERE app = ?
                ORDER BY ts_utc DESC
                LIMIT 200
                """,
                (app_name,),
            )

        source_search = st.text_input("Search article or document", key=f"{app_name}_history_search")
        search_rows = []
        if source_search:
            if _history_uses_sheets():
                needle = source_search.strip().lower()
                search_rows = [
                    {
                        "ts_utc": row["ts_utc"],
                        "email": row["email"],
                        "source_type": row["source_type"],
                        "source_label": row["source_label"],
                        "source_identity": row["source_identity"],
                        "iteration": row["iteration"],
                        "spelling_count": row["spelling_count"],
                        "grammar_count": row["grammar_count"],
                        "editorial_count": row["editorial_count"],
                        "fact_count": row["fact_count"],
                        "total_count": row["total_count"],
                    }
                    for row in all_analysis_rows
                    if needle in (row.get("source_label") or "").lower()
                    or needle in (row.get("source_identity") or "").lower()
                ][:200]
            else:
                like = f"%{source_search.strip()}%"
                search_rows = _fetch_rows(
                    """
                    SELECT
                        ts_utc,
                        email,
                        source_type,
                        source_label,
                        source_identity,
                        iteration,
                        spelling_count,
                        grammar_count,
                        editorial_count,
                        fact_count,
                        total_count
                    FROM analysis_runs
                    WHERE app = ?
                      AND (source_label LIKE ? OR source_identity LIKE ?)
                    ORDER BY ts_utc DESC
                    LIMIT 200
                    """,
                    (app_name, like, like),
                )

        st.markdown("#### Daily Summary")
        if daily_rows:
            st.dataframe(daily_rows, use_container_width=True)
        else:
            st.info("No daily history available yet.")

        st.markdown("#### Recent Analysis Runs")
        if recent_rows:
            st.dataframe(recent_rows, use_container_width=True)
        else:
            st.info("No analysis runs recorded yet.")

        st.markdown("#### Recent Logins")
        if login_rows:
            st.dataframe(login_rows, use_container_width=True)
        else:
            st.info("No login history recorded yet.")

        st.markdown("#### Article / Document Iterations")
        if source_search:
            if search_rows:
                st.dataframe(search_rows, use_container_width=True)
            else:
                st.info("No matching article or document history found.")
        else:
            st.caption("Search by URL, filename, or document hash to see iteration history.")

def enforce_app_access(app_title: str, app_caption: str, app_name: str):
    _restore_persisted_session(app_name)
    if _email_access_granted():
        with st.sidebar:
            st.caption(f"Signed in as {st.session_state.get('_email_access_email', '')}")
            if st.button("Log out"):
                _revoke_persisted_session(app_name)
                _clear_email_access()
                st.rerun()
        return

    st.title(app_title)
    st.caption(app_caption)
    with st.form("email_access_login"):
        username = st.text_input("Work email username", placeholder="firstname.lastname")
        st.caption(f"Domain fixed as @{ALLOWED_EMAIL_DOMAIN}")
        submitted = st.form_submit_button("Continue", type="primary")

    st.caption(APP_ACCESS_SUPPORT_TEXT)

    if submitted:
        email = _build_email_from_username(username)
        if not _email_allowed(email):
            st.error("Please enter only your username, without '@' or the domain.")
        else:
            st.session_state["_email_access_granted"] = True
            st.session_state["_email_access_email"] = email
            _record_access_event(app_name, email)
            _create_persisted_session(app_name, email)
            st.rerun()
    st.stop()

if not IMPORT_ONLY:
    enforce_app_access(
        "🧪 Hindi Article QC Tool (Gemini 2.5)",
        "Hindi Spelling · Grammar · Editorial Safety · AI Review",
        "hindi_qc",
    )
    st.title("🧪 Hindi Article QC Tool (Gemini 2.5)")
    st.caption("Hindi Spelling · Grammar · Editorial Safety · AI Review")

# =================================================
# AUTH CONFIG
# =================================================
PROJECT_ID = str(_secret("VERTEX_PROJECT_ID", "")).strip()
REGION = "us-central1"
CRED_PATH = "/tmp/gcp_service_account.json"
RULES_PATH = os.path.join(os.path.dirname(__file__), "hindi_qc_rules.txt")
MODEL_FLASH = "gemini-2.5-flash"
CLOUD_PLATFORM_SCOPE = "https://www.googleapis.com/auth/cloud-platform"
SHEETS_SCOPE = "https://www.googleapis.com/auth/spreadsheets"
PROMPT_VERSION_HI = "2026-04-01-2"
PERSISTENT_CACHE_PATH_HI = os.path.join(
    os.path.dirname(__file__),
    ".hindi_ai_output_cache.json",
)

# =================================================
# GCP AUTH
# =================================================
def load_gcp_credentials():
    creds, project_id, _ = _get_scoped_service_account_credentials([CLOUD_PLATFORM_SCOPE])
    return creds, project_id

# =================================================
# DYNAMIC QC RULES (OPTIONAL)
# =================================================
def load_hindi_rules():
    try:
        with open(RULES_PATH, "r", encoding="utf-8") as f:
            return f.read().strip()
    except FileNotFoundError:
        return ""
    except Exception:
        return ""

def load_hindi_rule_pairs():
    pairs = []
    for line in load_hindi_rules().splitlines():
        raw = line.strip()
        if not raw or raw.startswith("#") or "->" not in raw:
            continue
        wrong, correct = [part.strip() for part in raw.split("->", 1)]
        if not wrong or not correct or wrong == correct:
            continue
        pairs.append((wrong, correct))
    return pairs

# =================================================
# MODEL INIT (PARALLEL TO ENGLISH)
# =================================================
@st.cache_resource
def init_vertex_and_model():
    creds, project_id = load_gcp_credentials()

    client = genai.Client(
        vertexai=True,
        project=project_id,
        location=REGION,
        credentials=creds,
    )

    # Warmup (exactly like English)
    try:
        client.models.generate_content(
            model=MODEL_FLASH,
            contents="Warmup",
            config=genai_types.GenerateContentConfig(
                temperature=0,
                topP=1,
                maxOutputTokens=8,
            ),
        )
    except Exception:
        pass

    return client

def build_generate_config(generation_config=None):
    cfg = generation_config or {}
    return genai_types.GenerateContentConfig(
        temperature=cfg.get("temperature"),
        topP=cfg.get("top_p", 0),
        topK=cfg.get("top_k"),
        candidateCount=cfg.get("candidate_count"),
        maxOutputTokens=cfg.get("max_output_tokens"),
        seed=0,
        responseMimeType="text/plain",
        thinkingConfig=genai_types.ThinkingConfig(thinkingBudget=0),
    )

def generate_text(prompt, generation_config=None, model_name=MODEL_FLASH):
    client = init_vertex_and_model()
    response = client.models.generate_content(
        model=model_name,
        contents=prompt,
        config=build_generate_config(generation_config),
    )
    return response.text or ""

def format_ai_error(prefix: str, exc: Exception) -> str:
    return f"__ERROR__:{prefix}: {type(exc).__name__}: {exc}"

def is_ai_error_output(text: str) -> bool:
    return (text or "").startswith("__ERROR__:")

def snapshot_has_meaningful_output(snapshot: dict) -> bool:
    if not isinstance(snapshot, dict):
        return False
    return any((snapshot.get(key) or "").strip() for key in ("grammar_raw", "editorial_raw", "fact_result"))

NAVIGATION_TOKENS = {
    "अन्य", "मनोरंजन", "लाइफस्टाइल", "टेक-ज्ञान", "ऑटो", "पॉलिटिक्स",
    "did you know", "एक्सप्लेनर", "लाइव न्यूज़", "लाइव न्यूज़", "शिक्षा",
    "जॉब्स", "कैरियर", "वायरल", "स्पेशल", "वेब स्टोरी", "जागरण इमर्सिव",
}

ARTICLE_ROOT_SELECTORS = [
    "article",
    "[itemprop='articleBody']",
    ".article-content",
    ".article-body",
    ".articleBody",
    ".story-content",
    ".storyBody",
    ".entry-content",
    ".post-content",
    ".content-text",
    ".detail-content",
    ".description",
]

DOMAIN_ARTICLE_SELECTORS = {
    "jagran.com": [
        "article",
        ".ArticleDetail_ArticleDetail__NQJvJ",
        ".jg_m-article",
        "[itemprop='articleBody']",
    ],
    "www.jagran.com": [
        "article",
        ".ArticleDetail_ArticleDetail__NQJvJ",
        ".jg_m-article",
        "[itemprop='articleBody']",
    ],
    "herzindagi.com": [
        "[itemprop='articleBody']",
        "article",
        ".article-content",
        ".story-content",
        ".entry-content",
        ".post-content",
        ".description",
    ],
    "www.herzindagi.com": [
        "[itemprop='articleBody']",
        "article",
        ".article-content",
        ".story-content",
        ".entry-content",
        ".post-content",
        ".description",
    ],
    "onlymyhealth.com": [
        "[itemprop='articleBody']",
        "article",
        ".article-content",
        ".story-content",
        ".entry-content",
        ".post-content",
        ".description",
    ],
    "www.onlymyhealth.com": [
        "[itemprop='articleBody']",
        "article",
        ".article-content",
        ".story-content",
        ".entry-content",
        ".post-content",
        ".description",
    ],
}

EXCLUDED_SUBTREE_SELECTORS = [
    "script",
    "style",
    "noscript",
    "form",
    "nav",
    "figure",
    "figcaption",
    "footer",
    "aside",
    "header",
    ".breadcrumb",
    ".breadcrumbs",
    ".social-share",
    ".share",
    ".author",
    "[class*='author']",
    "[class*='Author']",
    ".byline",
    "[class*='byline']",
    ".updated",
    ".publish-info",
    ".highlights",
    ".highlight",
    ".keypoints",
    ".key-points",
    ".related",
    ".recommended",
    ".read-more",
    ".also-read",
    ".also_read",
    "[class*='shortInnerBlock']",
    ".you-may-like",
    ".trending",
    ".copyright",
    ".footer",
    ".ad",
    ".ads",
]

def is_navigation_blob(text: str) -> bool:
    compact = re.sub(r"\s+", " ", text.strip().lower())
    if len(compact) > 180:
        return False

    token_hits = sum(1 for token in NAVIGATION_TOKENS if token in compact)
    return token_hits >= 5

def is_probable_metadata_line(text: str) -> bool:
    compact = re.sub(r"\s+", " ", (text or "").strip())
    lower = compact.lower()

    if not compact:
        return True
    if re.match(r"^https?://", compact):
        return True
    if "www." in lower and lower.endswith(".html"):
        return True
    if lower in {"highlights", "highlight", "हाइलाइट्स"}:
        return True
    if lower.startswith("by ") or lower.startswith("edited by"):
        return True
    if "edited by:" in lower or lower.startswith("updated:") or lower.startswith("published:"):
        return True
    if "all rights reserved" in lower or "copyright" in lower:
        return True
    if "jagran new media" in lower:
        return True
    return False

def should_skip_extracted_text(text: str) -> bool:
    compact = re.sub(r"\s+", " ", text.strip())
    lower = compact.lower()

    if not compact:
        return True
    if is_probable_metadata_line(compact):
        return True
    if lower.startswith("यह भी पढ़ें"):
        return True
    if lower.startswith("...और पढ़ें") or lower.startswith("और पढ़ें"):
        return True
    if lower.startswith("जानिए मुख्य बातें") or "खबर का सार एक नजर में" in lower:
        return True
    if is_navigation_blob(compact):
        return True

    return False

def sanitize_extracted_text(text: str) -> str:
    cleaned = re.sub(r"\s+", " ", (text or "").strip())
    cleaned = re.sub(r"\s*\.{0,3}\s*और पढ़ें\s*$", "", cleaned)
    cleaned = re.sub(r"\s*यह भी पढ़ें[:：].*$", "", cleaned)
    cleaned = re.sub(r"\s*copyright\s+.*$", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\s*all rights reserved\s*.*$", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    return cleaned

HOUSE_STYLE_NUKTA_REPLACEMENTS = {
    "ज़्यादा": "ज्यादा",
    "ज़रूरी": "जरूरी",
    "बाज़ार": "बाजार",
    "नज़र": "नजर",
    "नज़रअंदाज़": "नजरअंदाज",
    "नज़रअंदाज": "नजरअंदाज",
    "नजरअंदाज़": "नजरअंदाज",
}

IYE_ENDING_EXCLUSIONS = {
    "लिए", "दिए", "किए", "लिये", "दिये", "किये",
    "जिए", "जीए", "पीए", "चाहिए", "आइए", "जाइए",
    "कहिए", "रहिए", "लीजिए", "दीजिए", "कीजिए", "पीजिए",
    "देखिए", "सुनिए", "मानिए", "जानिए", "बताइए", "अपनाइए",
    "हटाइए", "लगाइए", "करिए", "भरिए", "धरिए", "रखिए",
    "आईए", "जाईए",
}

IYE_ENDING_PATTERN = re.compile(
    r"(?<![A-Za-z0-9\u0900-\u097F])([A-Za-z\u0900-\u097F]{4,}िए)(?![A-Za-z0-9\u0900-\u097F])"
)

def apply_house_style_text_sanitizer(text: str) -> str:
    sanitized = text or ""
    sanitized = sanitized.replace("ँ", "ं")
    for wrong, correct in HOUSE_STYLE_NUKTA_REPLACEMENTS.items():
        sanitized = re.sub(
            rf"(?<![A-Za-z0-9\u0900-\u097F]){re.escape(wrong)}(?![A-Za-z0-9\u0900-\u097F])",
            correct,
            sanitized,
        )
    return sanitized

def infer_iye_ending_rows(article_data):
    rows = []
    seen = set()
    excluded = {normalise_hi(word) for word in IYE_ENDING_EXCLUSIONS}

    for ctype, text in article_data or []:
        if ctype not in {"heading", "paragraph", "table"}:
            continue

        sentences = split_hindi_sentences(text) or [text]
        for sentence in sentences:
            for match in IYE_ENDING_PATTERN.finditer(sentence):
                token = match.group(1)
                normalized = normalise_hi(token)
                if normalized in excluded:
                    continue
                corrected_token = re.sub(r"िए$", "िये", token)
                if corrected_token == token:
                    continue
                corrected_sentence = sentence.replace(token, corrected_token, 1)
                reason = f"वर्तनी / house style: use '{corrected_token}'"
                key = (canon_hi(sentence), canon_hi(corrected_sentence), canon_hi(reason))
                if key in seen:
                    continue
                seen.add(key)
                rows.append((sentence.strip(), corrected_sentence.strip(), reason))

    return rows

def get_domain(url: str) -> str:
    return (urlparse(url).netloc or "").lower()

def is_jagran_domain(url: str) -> bool:
    return get_domain(url) in {"jagran.com", "www.jagran.com"}

def has_inline_read_more(raw_text: str) -> bool:
    compact = re.sub(r"\s+", " ", (raw_text or "").strip())
    return "...और पढ़ें" in compact or compact.endswith("और पढ़ें")

def add_meta_description_summary(soup, url, content, seen):
    if not is_jagran_domain(url):
        return

    meta = (
        soup.find("meta", attrs={"name": "description"})
        or soup.find("meta", attrs={"property": "og:description"})
    )
    if not meta:
        return

    summary = sanitize_extracted_text(meta.get("content", ""))
    if len(summary) < 80:
        return
    if should_skip_extracted_text(summary):
        return
    if summary in seen:
        return

    seen.add(summary)
    content.append(("paragraph", summary))

def build_source_style_notes(source_context: str) -> str:
    domain = get_domain(source_context) if source_context else ""
    if domain in {
        "herzindagi.com",
        "www.herzindagi.com",
        "onlymyhealth.com",
        "www.onlymyhealth.com",
    }:
        return """
Domain style notes:
- In Hindi body copy for lifestyle, home, recipe, and health content, prefer clear Devanagari transliteration for standalone English common nouns or room/object labels when the term is not a brand name and the correction is unambiguous (for example, "Sink" -> "सिंक", "Living Room" -> "लिविंग रूम").
- Normalize broken Unicode half-letter / ZWJ rendering artifacts in Devanagari spellings (for example, forms like "स्‍टाइलिश", "अच्‍छी", "जिन्‍हें") to their standard rendered spellings.
- Use house-style joined compounds where the standard closed form is clear (for example, "रसोईघर").
- Use a hyphen in letter-shape compounds where appropriate (for example, "एल-शेप", "यू-शेप"), but do not add a hyphen to idiomatic expressions where it is not standard.
"""
    return ""

def is_sufficient_article_body(content) -> bool:
    paragraphs = [text for ctype, text in content if ctype == "paragraph"]
    total_chars = sum(len(text) for text in paragraphs)
    return len(paragraphs) >= 2 and total_chars >= 250

def extend_content_from_container(container, content, seen):
    clone = BeautifulSoup(str(container), "html.parser")
    for node in clone.select(",".join(EXCLUDED_SUBTREE_SELECTORS)):
        node.decompose()

    extracted_any = False
    seen_norms = {normalize_for_match(item) for item in seen}

    def should_add_text_candidate(txt: str) -> bool:
        norm_txt = normalize_for_match(txt)
        if not norm_txt:
            return False
        for existing_norm in seen_norms:
            if not existing_norm:
                continue
            if norm_txt == existing_norm or norm_txt in existing_norm or existing_norm in norm_txt:
                return False
        return True

    def append_text(ctype: str, txt: str):
        nonlocal extracted_any
        seen.add(txt)
        seen_norms.add(normalize_for_match(txt))
        content.append((ctype, txt))
        extracted_any = True

    for el in clone.find_all(["h2", "h3", "h4", "h5", "h6"], recursive=True):
        raw_txt = el.get_text(separator=" ", strip=True)
        if has_inline_read_more(raw_txt):
            continue
        txt = sanitize_extracted_text(raw_txt)
        if len(txt) < 8:
            continue
        if should_skip_extracted_text(txt):
            continue
        if txt in seen or not should_add_text_candidate(txt):
            continue
        append_text("heading", txt)

    for el in clone.find_all(["p", "li"], recursive=True):
        raw_txt = el.get_text(separator=" ", strip=True)
        if has_inline_read_more(raw_txt):
            continue
        txt = raw_txt
        txt = re.sub(r"\s+([,.;:!?])", r"\1", txt)
        txt = sanitize_extracted_text(txt)
        if not txt or len(txt) < 20:
            continue
        if should_skip_extracted_text(txt):
            continue
        if txt in seen or not should_add_text_candidate(txt):
            continue
        append_text("paragraph", txt)

    for table in clone.find_all("table"):
        for tr in table.find_all("tr"):
            cells = [
                c.get_text(separator=" ", strip=True)
                for c in tr.find_all(["th", "td"])
            ]
            cells = [re.sub(r"\s+", " ", c).strip() for c in cells if c.strip()]
            if not cells:
                continue
            row_text = sanitize_extracted_text(" | ".join(cells))
            if len(row_text) < 5:
                continue
            if should_skip_extracted_text(row_text):
                continue
            if row_text in seen or not should_add_text_candidate(row_text):
                continue
            append_text("table", row_text)

    fallback_text = clone.get_text(separator="\n", strip=True)
    for para in re.split(r"\n+", fallback_text):
        para = sanitize_extracted_text(para)
        min_len = 8 if is_heading_like_hi(para) else 20
        if len(para) < min_len:
            continue
        if should_skip_extracted_text(para):
            continue
        if para in seen or not should_add_text_candidate(para):
            continue
        append_text("heading" if is_heading_like_hi(para) else "paragraph", para)

def extract_from_article_roots(soup, url, content, seen):
    roots = []
    selectors = DOMAIN_ARTICLE_SELECTORS.get(get_domain(url), []) + ARTICLE_ROOT_SELECTORS

    for selector in selectors:
        for node in soup.select(selector):
            text = sanitize_extracted_text(node.get_text(separator=" ", strip=True))
            if len(text) < 150:
                continue
            roots.append(node)

    if not roots:
        return

    ordered = []
    for node in roots:
        if any(node in existing.descendants for existing in ordered):
            continue
        ordered.append(node)

    for node in ordered:
        extend_content_from_container(node, content, seen)

LDJSON_TEXT_FIELD_PATTERN = re.compile(
    r'"(?P<key>articleBody|text|description)"\s*:\s*"(?P<value>(?:[^"\\]|\\.|[\r\n])*)"',
    re.DOTALL,
)

def decode_ldjson_string(value: str) -> str:
    text = value or ""
    text = text.replace("\\n", "\n").replace("\\r", "\n").replace("\\t", " ")
    text = text.replace('\\"', '"').replace("\\/", "/").replace("\\\\", "\\")
    return text

def extract_text_fields_from_ldjson_raw(raw_script: str):
    fields = []
    if not raw_script:
        return fields

    for match in LDJSON_TEXT_FIELD_PATTERN.finditer(raw_script):
        key = match.group("key")
        value = decode_ldjson_string(match.group("value"))
        if len(value.strip()) > 80:
            fields.append((key, value))
    return fields

def extract_from_json_article_body(soup, content, seen):
    article_body_texts = []
    auxiliary_texts = []

    def extract_article_body(obj):
        if isinstance(obj, dict):
            for k, v in obj.items():
                if k in {"articleBody", "text", "description"} and isinstance(v, str):
                    if len(v) > 80:
                        if k == "articleBody":
                            article_body_texts.append(v)
                        else:
                            auxiliary_texts.append(v)
                else:
                    extract_article_body(v)
        elif isinstance(obj, list):
            for item in obj:
                extract_article_body(item)

    for script in soup.find_all("script", {"type": "application/ld+json"}):
        raw_script = script.string or script.get_text() or ""
        if not raw_script:
            continue
        try:
            data = json.loads(raw_script)
            extract_article_body(data)
        except Exception:
            for key, value in extract_text_fields_from_ldjson_raw(raw_script):
                if key == "articleBody":
                    article_body_texts.append(value)
                elif key in {"text", "description"}:
                    auxiliary_texts.append(value)

    body_texts = article_body_texts or auxiliary_texts

    for body in body_texts:
        cleaned = BeautifulSoup(body, "html.parser").get_text(separator="\n", strip=True)
        for para in re.split(r"\n+|\\n+", cleaned):
            para = sanitize_extracted_text(para)
            min_len = 8 if is_heading_like_hi(para) else 20
            if len(para) < min_len:
                continue
            if should_skip_extracted_text(para):
                continue
            if para in seen:
                continue
            seen.add(para)
            content.append(("heading" if is_heading_like_hi(para) else "paragraph", para))

# =================================================
# INPUT EXTRACTION (UNCHANGED STRUCTURE)
# =================================================
def clean_docx(file_path):
    from docx import Document
    doc = Document(file_path)

    content = []
    seen = set()

    for para in doc.paragraphs:
        txt = sanitize_extracted_text(para.text)
        if not txt or len(txt) < 15:
            continue
        if should_skip_extracted_text(txt):
            continue
        if txt in seen:
            continue

        seen.add(txt)
        content.append(("paragraph", txt))

    # Extract table rows (if any)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if not cells:
                continue
            row_text = sanitize_extracted_text(" | ".join(cells))
            if len(row_text) < 5:
                continue
            if should_skip_extracted_text(row_text):
                continue
            if row_text in seen:
                continue
            seen.add(row_text)
            content.append(("table", row_text))

    return content


def clean_article(url):
    headers = {"User-Agent": "Mozilla/5.0"}
    response = requests.get(url, headers=headers, timeout=15)
    response.raise_for_status()

    soup = BeautifulSoup(response.text, "html.parser")

    content = []
    seen = set()

    h1 = soup.find("h1")
    if h1:
        h1_text = sanitize_extracted_text(h1.get_text(strip=True))
        if h1_text:
            content.append(("heading", h1_text))
            seen.add(h1_text)

    if is_jagran_domain(url):
        extract_from_json_article_body(soup, content, seen)
        if is_sufficient_article_body(content):
            return content

    add_meta_description_summary(soup, url, content, seen)

    extract_from_article_roots(soup, url, content, seen)
    extract_from_json_article_body(soup, content, seen)

    if is_sufficient_article_body(content):
        return content

    for el in soup.find_all(["p", "li"], recursive=True):
        raw_txt = el.get_text(separator=" ", strip=True)
        if has_inline_read_more(raw_txt):
            continue
        txt = raw_txt
        txt = re.sub(r"\s+([,.;:!?])", r"\1", txt)
        txt = sanitize_extracted_text(txt)

        if not txt or len(txt) < 20:
            continue
        if should_skip_extracted_text(txt):
            continue
        if txt in seen:
            continue

        seen.add(txt)
        content.append(("paragraph", txt))

    for table in soup.find_all("table"):
        for tr in table.find_all("tr"):
            cells = [
                c.get_text(separator=" ", strip=True)
                for c in tr.find_all(["th", "td"])
            ]
            cells = [re.sub(r"\s+", " ", c).strip() for c in cells if c.strip()]
            if not cells:
                continue
            row_text = sanitize_extracted_text(" | ".join(cells))
            if len(row_text) < 5:
                continue
            if should_skip_extracted_text(row_text):
                continue
            if row_text in seen:
                continue
            seen.add(row_text)
            content.append(("table", row_text))

    if len(content) < 3:
        extract_from_json_article_body(soup, content, seen)

    return content

# =================================================
# STRUCTURAL LINE CHECK (PARITY WITH ENGLISH)
# =================================================
def is_structural_line_hi(text: str) -> bool:
    t = text.strip()

    if len(t.split()) <= 3:
        return True

    if re.match(r"^\d+[\).]", t):
        return True

    if any(t.startswith(x) for x in [
        "दिन ",
        "नोट:",
        "जानें",
        "देखें",
    ]):
        return True

    return False

# =================================================
# SENTENCE SPLITTER (spaCy replacement)
# =================================================
def split_hindi_sentences(text: str):
    parts = re.split(r"(।|\?|!)", text)
    sentences = []
    buf = ""

    for p in parts:
        buf += p
        if p in {"।", "?", "!"}:
            s = buf.strip()
            if len(s.split()) >= 6:
                sentences.append(s)
            buf = ""

    if buf.strip():
        sentences.append(buf.strip())

    return sentences

# =================================================
# CANONICAL NORMALISERS (EXPLICIT, LIKE ENGLISH)
# =================================================
def canon_hi(text: str) -> str:
    t = text.lower().strip()
    t = re.sub(r"\s+", " ", t)
    t = re.sub(r"[।,;:!?]", "", t)
    return t

def normalise_hi(text: str) -> str:
    return re.sub(r"[^\w\u0900-\u097F]", "", text.lower())

def normalize_for_match(text: str) -> str:
    if not text:
        return ""
    t = text.lower()
    t = re.sub(r"\s+", " ", t)
    t = t.replace("।", " ").replace(".", " ")
    t = re.sub(r"[\"'“”‘’]", "", t)
    t = re.sub(r"[^\w\u0900-\u097F ]", "", t)
    t = re.sub(r"\s+", " ", t).strip()
    return t

def normalize_quote_style(text: str) -> str:
    return (text or "").translate(str.maketrans({
        "“": '"',
        "”": '"',
        "„": '"',
        "‟": '"',
        "‘": "'",
        "’": "'",
        "‚": "'",
        "‛": "'",
    }))

NUKTA_STRIP_TRANSLATION = str.maketrans({
    ord("क़"): "क",
    ord("ख़"): "ख",
    ord("ग़"): "ग",
    ord("ज़"): "ज",
    ord("ड़"): "ड",
    ord("ढ़"): "ढ",
    ord("फ़"): "फ",
    ord("य़"): "य",
    ord("ऩ"): "न",
    ord("ऱ"): "र",
    ord("ऴ"): "ळ",
})

def strip_quote_chars(text: str) -> str:
    return re.sub(r"""["'“”‘’„‟‚‛]""", "", text or "")

def strip_nukta_chars(text: str) -> str:
    cleaned = (text or "").translate(NUKTA_STRIP_TRANSLATION)
    return cleaned.replace("\u093c", "")

def normalize_for_equality(text: str) -> str:
    return re.sub(r"\s+", " ", normalize_quote_style((text or "").strip()))

def is_noop_reason(reason: str) -> bool:
    lower = (reason or "").strip().lower()
    return lower in {
        "no corrections needed",
        "no correction needed",
        "no error",
        "no errors",
        "कोई त्रुटि नहीं",
        "कोई गलती नहीं",
    }

def is_noop_correction(original: str, corrected: str) -> bool:
    return normalize_for_equality(original) == normalize_for_equality(corrected)

def is_quote_only_correction(original: str, corrected: str) -> bool:
    original_base = normalize_for_equality(strip_quote_chars(original))
    corrected_base = normalize_for_equality(strip_quote_chars(corrected))
    if not original_base or not corrected_base:
        return False
    return original_base == corrected_base

def is_nukta_only_correction(original: str, corrected: str, reason: str) -> bool:
    if normalize_for_equality(original) == normalize_for_equality(corrected):
        return False
    original_base = normalize_for_equality(strip_nukta_chars(original))
    corrected_base = normalize_for_equality(strip_nukta_chars(corrected))
    if not original_base or not corrected_base or original_base != corrected_base:
        return False

    reason_lower = (reason or "").strip().lower()
    return any(marker in reason_lower for marker in (
        "वर्तनी",
        "spelling",
        "orthography",
        "loanword",
        "transliteration",
        "style",
    ))

def strip_punctuation_spacing(text: str) -> str:
    cleaned = normalize_quote_style((text or "").strip())
    cleaned = re.sub(r"\s*([,;:!?।])\s*", r"\1", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned

def is_heading_like_hi(text: str) -> bool:
    t = re.sub(r"\s+", " ", (text or "").strip())
    if not t:
        return False
    if len(t) > 140:
        return False

    word_count = len(t.split())
    if word_count <= 12 and not re.search(r"[।!?]$", t):
        return True

    if ":" in t and word_count <= 18:
        return True

    if word_count <= 12 and t.endswith("?") and re.search(r"(क्या|कैसे|क्यों|कब|कितना|कौन|किस)", t):
        return True

    heading_markers = (
        "fact check:",
        "फैक्ट चेक:",
        "exclusive:",
        "explained:",
        "live:",
        "photo gallery:",
        "video:",
    )
    lower = t.lower()
    return any(lower.startswith(marker) for marker in heading_markers)

def is_heading_danda_correction(original: str, corrected: str, reason: str) -> bool:
    original_norm = normalize_for_equality(original)
    corrected_norm = normalize_for_equality(corrected)
    if not original_norm or not corrected_norm:
        return False

    reason_lower = (reason or "").strip().lower()
    if not any(marker in reason_lower for marker in (
        "पूर्ण विराम",
        "sentence-ending punctuation",
        "danda",
        "punctuation",
    )):
        return False

    if not is_heading_like_hi(original_norm):
        return False

    return corrected_norm == f"{original_norm}।"

def is_ignored_styleguide_issue(text: str) -> bool:
    lower = (text or "").strip().lower()
    markers = (
        "prefer chandra-bindu",
        "use chandra-bindu",
        "चंद्रबिंदु का प्रयोग करें",
        "चन्द्रबिन्दु का प्रयोग करें",
    )
    return any(marker in lower for marker in markers)

def is_self_contradictory_reason(original: str, corrected: str, reason: str) -> bool:
    reason_text = reason or ""
    if not reason_text:
        return False

    quote_pairs = re.findall(
        r"""['"“”‘’]([^'"“”‘’]+)['"“”‘’]\s*(?:को|to)\s*['"“”‘’]([^'"“”‘’]+)['"“”‘’]""",
        reason_text,
        flags=re.IGNORECASE,
    )
    for source, target in quote_pairs:
        if normalize_for_equality(source) == normalize_for_equality(target):
            return True

    return False

def is_bad_punctuation_spacing_correction(original: str, corrected: str, reason: str) -> bool:
    reason_lower = (reason or "").strip().lower()
    if not any(marker in reason_lower for marker in (
        "विराम",
        "punctuation",
        "spacing",
        "space",
        "comma",
    )):
        return False

    if strip_punctuation_spacing(original) != strip_punctuation_spacing(corrected):
        return False

    original_text = normalize_quote_style(original or "")
    corrected_text = normalize_quote_style(corrected or "")
    had_post_punct_space = re.search(r"[,;:!?।]\s+[A-Za-z\u0900-\u097F]", original_text)
    lost_post_punct_space = re.search(r"[,;:!?।][A-Za-z\u0900-\u097F]", corrected_text)
    return bool(had_post_punct_space and lost_post_punct_space)

def is_redundant_gender_rewrite(original: str, corrected: str, reason: str) -> bool:
    reason_lower = (reason or "").strip().lower()
    if not any(marker in reason_lower for marker in (
        "वर्तनी",
        "spelling",
        "grammar",
        "orthography",
        "house style",
        "style",
    )):
        return False

    original_tokens = (original or "").split()
    corrected_tokens = (corrected or "").split()
    if len(original_tokens) != len(corrected_tokens) or len(original_tokens) < 2:
        return False

    feminine_contexts = {
        "महिला",
        "स्त्री",
        "लड़की",
        "युवती",
        "बालिका",
    }
    if normalise_hi(original_tokens[1]) not in {normalise_hi(x) for x in feminine_contexts}:
        return False
    if normalise_hi(original_tokens[1]) != normalise_hi(corrected_tokens[1]):
        return False

    if any(
        normalise_hi(o) != normalise_hi(c)
        for o, c in zip(original_tokens[2:], corrected_tokens[2:])
    ):
        return False

    original_head = normalise_hi(original_tokens[0])
    corrected_head = normalise_hi(corrected_tokens[0])
    if not original_head or not corrected_head or original_head == corrected_head:
        return False

    shared_prefix = os.path.commonprefix([original_head, corrected_head])
    if len(shared_prefix) < max(3, min(len(original_head), len(corrected_head)) - 1):
        return False

    return True

def is_ambiguous_homophone_correction(original: str, corrected: str, reason: str) -> bool:
    original_tokens = word_tokens_hi(original)
    corrected_tokens = word_tokens_hi(corrected)
    if len(original_tokens) != len(corrected_tokens) or not original_tokens:
        return False

    diffs = []
    for original_token, corrected_token in zip(original_tokens, corrected_tokens):
        if normalise_hi(original_token) != normalise_hi(corrected_token):
            diffs.append((original_token, corrected_token))

    if len(diffs) != 1:
        return False

    original_token, corrected_token = diffs[0]
    if normalise_hi(original_token) != normalise_hi("काफी") or normalise_hi(corrected_token) != normalise_hi("कॉफी"):
        return False

    beverage_markers = {
        "पीना", "पिएं", "पियो", "कॉफी", "कप", "मग", "कैफे", "कैफीन",
        "ब्रू", "दूध", "चीनी", "पेय", "कैप्पुचीनो", "एस्प्रेसो",
    }
    context = f"{original} {corrected} {reason}".lower()
    return not any(marker in context for marker in beverage_markers)

def should_skip_language_change(original: str, corrected: str, reason: str) -> bool:
    return any((
        is_noop_reason(reason),
        is_noop_correction(original, corrected),
        is_quote_only_correction(original, corrected),
        is_nukta_only_correction(original, corrected, reason),
        is_heading_danda_correction(original, corrected, reason),
        is_ignored_styleguide_issue(reason),
        is_self_contradictory_reason(original, corrected, reason),
        is_bad_punctuation_spacing_correction(original, corrected, reason),
        is_redundant_gender_rewrite(original, corrected, reason),
        is_ambiguous_homophone_correction(original, corrected, reason),
    ))

def should_project_editorial_to_language(issue: str) -> bool:
    lower = (issue or "").strip().lower()
    if not lower:
        return False
    return any(token in lower for token in (
        "spelling",
        "typo",
        "grammar",
        "punctuation",
        "spacing",
        "quotation",
        "quote style for direct speech",
        "sentence ending punctuation",
    ))

def is_no_issue_fact(issue: str, correction: str) -> bool:
    issue_lower = (issue or "").strip().lower()
    correction_lower = (correction or "").strip().lower()
    if issue_lower in {"", "-", "--", "---", "no issue", "no issues"}:
        return True
    if correction_lower in {"", "-", "--", "---", "no issue", "no issues"}:
        return True
    return False

def is_style_only_fact(statement: str, issue: str, correction: str) -> bool:
    lower_issue = (issue or "").strip().lower()
    lower_correction = (correction or "").strip().lower()
    combined = f"{lower_issue} {lower_correction}"

    if canon_hi(statement) == canon_hi(correction):
        return True

    style_markers = (
        "spelling",
        "grammar",
        "punctuation",
        "style",
        "format",
        "wording",
        "terminology",
        "abbreviation",
        "full form",
        "quote",
        "comma",
        "should be",
        "use ",
        "replace ",
        "preferred",
    )
    return any(marker in combined for marker in style_markers)

def find_context_snippet(article_data, needle: str) -> str:
    if not needle:
        return ""

    norm_needle = normalize_for_match(needle)
    for ctype, text in article_data:
        if ctype not in {"paragraph", "table"}:
            continue
        if norm_needle not in normalize_for_match(text):
            continue

        sentences = split_hindi_sentences(text)
        if not sentences:
            sentences = [text]

        for sentence in sentences:
            if norm_needle in normalize_for_match(sentence):
                return sentence.strip()

        return text.strip()

    return ""

def needs_context_expansion(original: str, reason: str) -> bool:
    if not original:
        return False

    token_count = len(original.split())
    lower_reason = (reason or "").lower()
    if token_count <= 2:
        return True
    if len(original) < 18:
        return True
    if "abbreviation" in lower_reason or "full form" in lower_reason:
        return True

    return False

def expand_language_row_context(article_data, original: str, corrected: str, reason: str):
    if not needs_context_expansion(original, reason):
        return original, corrected, reason

    context = find_context_snippet(article_data, original)
    if not context:
        return original, corrected, reason
    if canon_hi(context) == canon_hi(original):
        return original, corrected, reason
    if original not in context:
        return original, corrected, reason

    corrected_context = context.replace(original, corrected, 1)
    if canon_hi(corrected_context) == canon_hi(context):
        return original, corrected, reason

    return context, corrected_context, reason

def rule_based_spelling_rows(article_data):
    rows = []
    seen = set()
    for wrong, correct in load_hindi_rule_pairs():
        pattern = re.compile(
            rf"(?<![A-Za-z0-9\u0900-\u097F]){re.escape(wrong)}(?![A-Za-z0-9\u0900-\u097F])"
        )
        for ctype, text in article_data or []:
            if ctype not in {"heading", "paragraph", "table"}:
                continue
            if wrong not in text:
                continue

            sentences = split_hindi_sentences(text) or [text]
            for sentence in sentences:
                if not pattern.search(sentence):
                    continue

                corrected_sentence = pattern.sub(correct, sentence, count=1)
                reason = f"वर्तनी / house style: use '{correct}'"
                key = (canon_hi(sentence), canon_hi(corrected_sentence), canon_hi(reason))
                if key in seen:
                    continue
                seen.add(key)
                rows.append((sentence.strip(), corrected_sentence.strip(), reason))

    for sentence, corrected_sentence, reason in infer_iye_ending_rows(article_data):
        key = (canon_hi(sentence), canon_hi(corrected_sentence), canon_hi(reason))
        if key in seen:
            continue
        seen.add(key)
        rows.append((sentence, corrected_sentence, reason))

    return rows

def batch_hindi_texts(texts, max_chars=6000):
    batches = []
    current = []
    current_len = 0

    for text in texts:
        chunk_len = len(text) + 2
        if current and current_len + chunk_len > max_chars:
            batches.append("\n\n".join(current))
            current = [text]
            current_len = chunk_len
        else:
            current.append(text)
            current_len += chunk_len

    if current:
        batches.append("\n\n".join(current))

    return batches

def segment_hindi_review_text(text: str, max_chars: int = 520):
    cleaned = sanitize_extracted_text(text)
    if not cleaned:
        return []
    if len(cleaned) <= max_chars:
        return [cleaned]

    sentences = split_hindi_sentences(cleaned)
    if not sentences:
        sentences = [cleaned]

    segments = []
    current = []
    current_len = 0

    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue
        sentence_len = len(sentence) + (1 if current else 0)

        if current and current_len + sentence_len > max_chars:
            segments.append(" ".join(current).strip())
            current = [sentence]
            current_len = len(sentence)
            continue

        current.append(sentence)
        current_len += sentence_len

    if current:
        segments.append(" ".join(current).strip())

    if segments:
        return segments

    return [cleaned]

# =================================================
# VERBATIM ROW FILTER (UNCHANGED LOGIC)
# =================================================
def filter_gemini_rows(raw_table, article_text):
    lines = raw_table.splitlines()
    output = []
    header_added = False
    norm_article = normalize_for_match(article_text)

    for line in lines:
        if line.strip().startswith("| Original"):
            if not header_added:
                output.append("| Original | Corrected | Reason |")
                output.append("|---|---|---|")
                header_added = True
            continue

        if "|" not in line:
            continue

        cols = [c.strip() for c in line.split("|") if c.strip()]
        if len(cols) != 3:
            continue

        original, corrected, reason = cols

        lower_cols = [c.lower() for c in cols]
        if lower_cols == ["original", "corrected", "reason"]:
            continue
        if any(c in {"original", "corrected", "reason"} for c in lower_cols):
            continue
        if any(c.strip() in {"-", "--", "---"} for c in cols):
            continue
        if should_skip_language_change(original, corrected, reason):
            continue

        if normalize_for_match(original) in norm_article:
            if not header_added:
                output.append("| Original | Corrected | Reason |")
                output.append("|---|---|---|")
                header_added = True
            output.append(f"| {original} | {corrected} | {reason} |")

    return "\n".join(output) if header_added else ""

# =================================================
# SPELLING VS GRAMMAR CLASSIFIER (EXPLICIT LAYER)
# =================================================
def is_hindi_spelling_issue(original, corrected):
    return (
        len(original.split()) == 1 and
        len(corrected.split()) == 1 and
        original != corrected
    )

def word_tokens_hi(text: str):
    return re.findall(r"[A-Za-z0-9\u0900-\u097F]+", text or "")

def changed_word_token_count(original: str, corrected: str) -> int:
    original_tokens = [normalise_hi(token) for token in word_tokens_hi(original)]
    corrected_tokens = [normalise_hi(token) for token in word_tokens_hi(corrected)]
    matcher = SequenceMatcher(a=original_tokens, b=corrected_tokens)
    changed = 0
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            continue
        changed += max(i2 - i1, j2 - j1)
    return changed

def looks_like_sentence_level_spelling_change(original: str, corrected: str) -> bool:
    original_tokens = word_tokens_hi(original)
    corrected_tokens = word_tokens_hi(corrected)
    if not original_tokens or not corrected_tokens:
        return False
    if abs(len(original_tokens) - len(corrected_tokens)) > 1:
        return False
    changed = changed_word_token_count(original, corrected)
    return 1 <= changed <= 2

def is_spelling_reason(reason: str) -> bool:
    lower = (reason or "").strip().lower()
    return any(marker in lower for marker in (
        "वर्तनी",
        "spelling",
        "typo",
        "misspelling",
        "orthography",
        "matra",
        "loanword",
        "transliteration",
        "चंद्रबिंदु",
        "चन्द्रबिन्दु",
        "anusvara",
        "chandrabindu",
        "house-style",
        "preferred spelling",
    ))

def classify_language_issue(original, corrected, reason):
    if is_spelling_reason(reason):
        return "spelling"
    if is_hindi_spelling_issue(original, corrected):
        return "spelling"
    if looks_like_sentence_level_spelling_change(original, corrected):
        return "spelling"
    return "grammar"

def split_grapheme_like_units(text: str):
    units = []
    current = ""

    for ch in text or "":
        if not current:
            current = ch
            continue

        if (
            unicodedata.combining(ch)
            or ch in {"\u200c", "\u200d", "\ufe0f", "\u094d"}
            or current.endswith(("\u094d", "\u200c", "\u200d"))
        ):
            current += ch
            continue

        units.append(current)
        current = ch

    if current:
        units.append(current)

    return units

def classify_diff_unit(unit: str) -> str:
    if not unit:
        return "other"
    if unit.isspace():
        return "space"

    first = unit[0]
    category = unicodedata.category(first)
    if category[0] in {"L", "M", "N"}:
        return "word"

    return "punct"

def tokenize_for_diff(text: str):
    tokens = []
    current = []
    current_kind = None

    for unit in split_grapheme_like_units(text or ""):
        kind = classify_diff_unit(unit)

        if kind == "punct":
            if current:
                tokens.append("".join(current))
                current = []
                current_kind = None
            tokens.append(unit)
            continue

        if current and kind != current_kind:
            tokens.append("".join(current))
            current = []

        current.append(unit)
        current_kind = kind

    if current:
        tokens.append("".join(current))

    return tokens

def highlight_diff_pair(original: str, corrected: str):
    original_tokens = tokenize_for_diff(original or "")
    corrected_tokens = tokenize_for_diff(corrected or "")
    matcher = SequenceMatcher(a=original_tokens, b=corrected_tokens)

    original_parts = []
    corrected_parts = []

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        original_chunk = "".join(original_tokens[i1:i2])
        corrected_chunk = "".join(corrected_tokens[j1:j2])

        if tag == "equal":
            original_parts.append(html.escape(original_chunk))
            corrected_parts.append(html.escape(corrected_chunk))
            continue

        if original_chunk:
            original_parts.append(
                f'<span class="qc-diff qc-diff-original">{html.escape(original_chunk)}</span>'
            )
        if corrected_chunk:
            corrected_parts.append(
                f'<span class="qc-diff qc-diff-corrected">{html.escape(corrected_chunk)}</span>'
            )

    return "".join(original_parts), "".join(corrected_parts)

def render_language_table(rows):
    if not rows:
        return ""

    lines = [
        """
<style>
.qc-table {
  width: 100%;
  border-collapse: collapse;
  margin-bottom: 1rem;
}
.qc-table th, .qc-table td {
  border: 1px solid rgba(250,250,250,0.14);
  padding: 0.75rem 0.9rem;
  vertical-align: top;
  text-align: left;
}
.qc-table .qc-diff-original {
  color: #ff6b6b;
  font-weight: 700;
}
.qc-table .qc-diff-corrected {
  color: #4ade80;
  font-weight: 700;
}
</style>
<table class="qc-table">
  <thead>
    <tr>
      <th>Original</th>
      <th>Corrected</th>
      <th>Reason</th>
    </tr>
  </thead>
  <tbody>
        """.strip()
    ]

    for original, corrected, reason in rows:
        original_html, corrected_html = highlight_diff_pair(original, corrected)
        lines.append(
            "<tr>"
            f"<td>{original_html}</td>"
            f"<td>{corrected_html}</td>"
            f"<td>{html.escape(reason)}</td>"
            "</tr>"
        )

    lines.append("</tbody></table>")
    return "\n".join(lines)

def parse_markdown_table_rows(table_md: str, expected_columns: int):
    rows = []
    for line in (table_md or "").splitlines():
        row = line.strip()
        if not row.startswith("|") or row.count("|") < expected_columns:
            continue
        parts = [part.strip() for part in row.strip("|").split("|")]
        if len(parts) != expected_columns:
            continue
        if all(re.fullmatch(r":?-{2,}:?", part or "") for part in parts):
            continue
        if parts[0].lower() in {"statement", "issue", "original"}:
            continue
        rows.append(parts)
    return rows

def find_hindi_pdf_font_path():
    candidates = [
        "/usr/share/fonts/truetype/noto/NotoSansDevanagari-Regular.ttf",
        "/usr/share/fonts/opentype/noto/NotoSansDevanagari-Regular.ttf",
        "/usr/share/fonts/truetype/noto/NotoSerifDevanagari-Regular.ttf",
        "/usr/share/fonts/opentype/noto/NotoSerifDevanagari-Regular.ttf",
    ]
    for path in candidates:
        if os.path.exists(path):
            return path
    return ""

def build_hindi_qc_report_pdf(source_label: str, user_email: str, spelling_rows, grammar_rows, editorial_rows, fact_md: str):
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    except Exception:
        return None, "PDF export requires the `reportlab` package."

    font_path = find_hindi_pdf_font_path()
    if not font_path:
        return None, "PDF export requires a Devanagari font on the server."

    font_name = "HindiReportFont"
    try:
        if font_name not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont(font_name, font_path))
    except Exception:
        return None, "PDF export could not load the Hindi font."

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "qc-title",
        parent=styles["Title"],
        fontName=font_name,
        fontSize=18,
        leading=22,
        textColor=colors.HexColor("#111827"),
    )
    body_style = ParagraphStyle(
        "qc-body",
        parent=styles["BodyText"],
        fontName=font_name,
        fontSize=9.5,
        leading=13,
        textColor=colors.HexColor("#111827"),
    )
    heading_style = ParagraphStyle(
        "qc-heading",
        parent=styles["Heading2"],
        fontName=font_name,
        fontSize=12,
        leading=16,
        textColor=colors.HexColor("#1f2937"),
        spaceBefore=10,
        spaceAfter=6,
    )

    def p(text):
        safe = html.escape((text or "").replace("\n", " "))
        return Paragraph(safe, body_style)

    def add_table(story, title, headers, rows, column_widths):
        story.append(Paragraph(title, heading_style))
        if not rows:
            story.append(Paragraph("No issues found", body_style))
            story.append(Spacer(1, 0.2 * cm))
            return

        table_data = [[Paragraph(html.escape(h), body_style) for h in headers]]
        for row in rows:
            table_data.append([p(cell) for cell in row])

        table = Table(table_data, colWidths=column_widths, repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E5E7EB")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#111827")),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#D1D5DB")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]))
        story.append(table)
        story.append(Spacer(1, 0.25 * cm))

    fact_rows = parse_markdown_table_rows(fact_md, 3)
    summary_rows = [
        ["Source", source_label or "-"],
        ["User", user_email or "-"],
        ["Generated (UTC)", _utc_now().strftime("%Y-%m-%d %H:%M:%S")],
        ["Spelling issues", str(len(spelling_rows or []))],
        ["Grammar issues", str(len(grammar_rows or []))],
        ["Editorial issues", str(len(editorial_rows or []))],
        ["Fact check issues", str(len(fact_rows))],
    ]

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=1.2 * cm,
        leftMargin=1.2 * cm,
        topMargin=1.2 * cm,
        bottomMargin=1.2 * cm,
    )

    story = [
        Paragraph("Hindi QC Report", title_style),
        Spacer(1, 0.25 * cm),
    ]

    add_table(story, "Summary", ["Field", "Value"], summary_rows, [4.2 * cm, 12.4 * cm])
    add_table(story, "Spelling Issues", ["Original", "Corrected", "Reason"], spelling_rows or [], [6.1 * cm, 6.1 * cm, 4.4 * cm])
    add_table(story, "Grammar Issues", ["Original", "Corrected", "Reason"], grammar_rows or [], [6.1 * cm, 6.1 * cm, 4.4 * cm])
    add_table(story, "Gemini Editorial Review", ["Issue", "Location", "Excerpt", "Corrected Text"], editorial_rows or [], [3.2 * cm, 2.4 * cm, 5.2 * cm, 6.0 * cm])
    add_table(story, "Fact Check", ["Statement", "Issue", "Correct Fact"], fact_rows, [6.0 * cm, 4.0 * cm, 6.8 * cm])

    doc.build(story)
    return buffer.getvalue(), None

def split_spelling_grammar_hi(table_md):
    spelling_rows = []
    grammar_rows = []

    rows = re.findall(
        r"\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|",
        table_md
    )

    for o, c, r in rows:
        if o.lower() == "original":
            continue

        reason = r.strip()
        if not reason or reason in {"-", "--", "---"}:
            continue
        if should_skip_language_change(o, c, reason):
            continue

        if classify_language_issue(o, c, reason) == "spelling":
            spelling_rows.append((o, c, reason))
        else:
            grammar_rows.append((o, c, reason))

    return spelling_rows, grammar_rows

def parse_language_rows(table_md, article_data=None):
    rows = []
    seen = set()

    matches = re.findall(
        r"\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|",
        table_md or ""
    )

    for original, corrected, reason in matches:
        if original.lower() == "original":
            continue
        if all(re.fullmatch(r":?-{2,}:?", cell.strip()) for cell in (original, corrected, reason)):
            continue
        if any(x.strip() in {"-", "--", "---"} for x in (original, corrected, reason)):
            continue
        if should_skip_language_change(original, corrected, reason):
            continue

        key = (canon_hi(original), canon_hi(corrected), canon_hi(reason))
        if key in seen:
            continue

        seen.add(key)
        original, corrected, reason = (
            original.strip(),
            corrected.strip(),
            reason.strip(),
        )
        corrected = apply_house_style_text_sanitizer(corrected)
        if article_data:
            original, corrected, reason = expand_language_row_context(
                article_data, original, corrected, reason
            )
            corrected = apply_house_style_text_sanitizer(corrected)
        if should_skip_language_change(original, corrected, reason):
            continue
        rows.append((original, corrected, reason))

    return rows

def parse_editorial_rows(editorial_md, article_data=None):
    rows = []
    seen = set()

    matches = re.findall(
        r"\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|",
        editorial_md or ""
    )

    for issue, location, excerpt, corrected in matches:
        lower = [issue.lower(), location.lower(), excerpt.lower(), corrected.lower()]
        if lower == ["issue", "location", "excerpt", "corrected text"]:
            continue
        if all(re.fullmatch(r":?-{2,}:?", cell.strip()) for cell in (issue, location, excerpt, corrected)):
            continue
        if issue.lower() == "no issues found":
            continue
        if any(x.strip() in {"-", "--", "---"} for x in (issue, location, excerpt, corrected)):
            continue

        issue, location, excerpt, corrected = (
            issue.strip(),
            location.strip(),
            excerpt.strip(),
            corrected.strip(),
        )
        corrected = apply_house_style_text_sanitizer(corrected)
        if should_skip_language_change(excerpt, corrected, issue):
            continue
        if article_data:
            excerpt, corrected, _ = expand_language_row_context(
                article_data, excerpt, corrected, issue
            )
            corrected = apply_house_style_text_sanitizer(corrected)
        if should_skip_language_change(excerpt, corrected, issue):
            continue

        key = (canon_hi(issue), canon_hi(location), canon_hi(excerpt), canon_hi(corrected))
        if key in seen:
            continue

        seen.add(key)
        rows.append((issue, location, excerpt, corrected))

    return rows

def parse_editorial_as_language_rows(editorial_md, article_data=None):
    rows = []
    seen = set()

    for issue, location, excerpt, corrected in parse_editorial_rows(editorial_md, article_data):
        if not should_project_editorial_to_language(issue):
            continue
        key = (canon_hi(excerpt), canon_hi(corrected), canon_hi(issue))
        if key in seen:
            continue

        seen.add(key)
        rows.append((excerpt, corrected, issue))

    return rows

def build_editorial_table(editorial_rows):
    if not editorial_rows:
        return ""

    lines = [
        "| Issue | Location | Excerpt | Corrected Text |",
        "|---|---|---|---|",
    ]
    for issue, location, excerpt, corrected in editorial_rows:
        lines.append(f"| {issue} | {location} | {excerpt} | {corrected} |")

    return "\n".join(lines)

def build_language_tables(language_rows, editorial_rows=None):
    spelling_rows = []
    grammar_rows = []
    seen = set()

    for original, corrected, reason in (language_rows or []) + (editorial_rows or []):
        if not original or not corrected or not reason:
            continue
        if should_skip_language_change(original, corrected, reason):
            continue

        key = (canon_hi(original), canon_hi(corrected), canon_hi(reason))
        if key in seen:
            continue

        seen.add(key)
        row = (original, corrected, reason)

        if classify_language_issue(original, corrected, reason) == "spelling":
            spelling_rows.append(row)
        else:
            grammar_rows.append(row)

    return spelling_rows, grammar_rows

# =================================================
# EDITORIAL ROW FILTER (HINDI)
# =================================================
def filter_editorial_rows(raw_table, article_text):
    lines = raw_table.splitlines()
    output = []
    header_added = False
    norm_article = normalize_for_match(article_text)

    for line in lines:
        if line.strip().startswith("| Issue"):
            if not header_added:
                output.append("| Issue | Location | Excerpt | Corrected Text |")
                output.append("|---|---|---|---|")
                header_added = True
            continue

        if "|" not in line:
            continue

        cols = [c.strip() for c in line.split("|") if c.strip()]
        if len(cols) != 4:
            continue

        issue, location, excerpt, corrected = cols

        lower_cols = [c.lower() for c in cols]
        if lower_cols == ["issue", "location", "excerpt", "corrected text"]:
            continue
        if any(c in {"issue", "location", "excerpt", "corrected text"} for c in lower_cols):
            continue
        if any(c.strip() in {"-", "--", "---"} for c in cols):
            continue

        # Excerpt must be present in source text
        if normalize_for_match(excerpt) not in norm_article:
            continue

        # Corrected must be different from excerpt (after trimming)
        if should_skip_language_change(excerpt, corrected, issue):
            continue

        if not header_added:
            output.append("| Issue | Location | Excerpt | Corrected Text |")
            output.append("|---|---|---|---|")
            header_added = True
        output.append(f"| {issue} | {location} | {excerpt} | {corrected} |")

    return "\n".join(output) if header_added else ""

# =================================================
# GEMINI GRAMMAR QC (PARAGRAPH PASS)
# =================================================
def gemini_grammar_review(article_data, source_context=""):
    raw_paragraphs = []
    for ctype, text in article_data:
        if ctype not in {"heading", "paragraph", "table"}:
            continue
        if ctype == "paragraph" and is_structural_line_hi(text):
            continue
        raw_paragraphs.extend(segment_hindi_review_text(text))

    raw_paragraphs = raw_paragraphs[:180]

    if not raw_paragraphs:
        return ""

    rules_text = load_hindi_rules()
    rules_block = (
        "\nOptional preferred spellings (hints only; do not limit your review to this list):\n"
        + rules_text + "\n"
        if rules_text else
        "\nOptional preferred spellings: (none provided)\n"
    )
    source_style_notes = build_source_style_notes(source_context)

    BASE_PROMPT = f"""
You are a professional Hindi editor and content QC reviewer.

Scope:
- Review the paragraph carefully from start to end
- Fix spelling, grammar, punctuation, and formatting errors in Hindi
- Do NOT translate, summarize, or rewrite
- Do NOT change names, places, numbers, or quotes unless the spelling is clearly wrong and the correction is unambiguous

Must-follow Hindi editorial rules:
- Use the Hindi danda "।" to end sentences (not a period).
- Do not add a danda to headlines, decks, labels, or subheadings that are not full sentences.
- Use double quotes for direct speech and official statements.
- Use single quotes for titles (books, films, shows, programs, named schemes).
- Straight and curly variants of the same quote type are both acceptable; do not flag quote-shape-only swaps such as "..." vs “...”.
- This publication style does not use chandrabindu in normal spellings where the house style prefers anusvara or the non-chandrabindu form; for example, prefer "पांच" over "पाँच".
- Do not suggest nukta-only rewrites in ordinary Hindi words when the non-nukta spelling is already acceptable house style (for example, do not force ज्यादा -> ज़्यादा, जरूरी -> ज़रूरी, बाजार -> बाज़ार, नजर -> नज़र).
- Do not introduce chandrabindu in corrected text unless it is unquestionably required by the publication style.
- For established loanwords that conventionally use the "ऑ" sound, prefer the standard spelling with "ऑ" when the correction is truly unambiguous (for example, "कापी" -> "कॉपी", "कालेज" -> "कॉलेज"). Do not change a normal Hindi word just because it resembles a loanword; for example, do not replace "काफी" meaning "enough/quite" with "कॉफी".
- For transliterated proper nouns or foreign names ending in "िए", prefer the house-style "िये" ending when that is clearly the intended pronunciation/transliteration (for example, "तुर्किए" -> "तुर्किये"). Do not apply this to ordinary Hindi verb forms such as "लिए", "दिए", "किए", "चाहिए", or "दीजिए".
- Flag first-mention abbreviation issues only when the short form refers to a named entity
  (such as an organisation, authority, institution, political party, law, scheme, or court)
  and the expansion is genuinely needed for clarity.
- Do not force expansion of common technical abbreviations, scientific labels, measurements,
  or interface tags.
- If a headline/subheading contains क्या/कैसे/क्यों/कब/कितना, it must end with "?".
- Use exactly three dots for ellipsis "...", not more.
- Use numbers 1–9 in words, 10+ in numerals (except dates, time, prices, recipe ingredients).
- Do not use honorifics श्री/श्रीमती for any person name (only "महामहिम" for the President when needed).

{rules_block}
{source_style_notes}

Guidance:
- The preferred spellings list is optional and incomplete. Use it as hints only.
- Still detect and flag other spelling/grammar issues dynamically.
- For word-level orthography, matra, chandrabindu, anusvara, or standard loanword-form corrections, treat the issue as a spelling issue and make the reason explicitly mention spelling/वर्तनी.
- Treat first-mention abbreviation problems as valid issues only when the short form is unclear
  without expansion in that article context.
- Do not stop after finding the first issue in a paragraph.
- Identify all clear issues in the paragraph, including quote misuse, punctuation, spacing, wording, and abbreviation-introduction problems.
- Apply the house orthography consistently: avoid chandrabindu-style spellings in normal words when the house style prefers non-chandrabindu forms, and preserve standard "ऑ" loanword spellings where clearly appropriate.
- Prefer the existing non-nukta house-style spellings for ordinary Hindi words unless a proper noun or an unquestionably fixed foreign spelling requires nukta.
- Do not enforce subjective style preferences such as replacing acceptable loanwords,
  banning sentence openings like "लेकिन", or mandating commas after specific discourse markers.
- Do not translate acceptable English technical terms into Hindi just to make a correction.

Constraints:
- Use only TEXT
- Original must be exact substring
- No hallucination

Return output strictly as table with header:
| Original | Corrected | Reason |

Reason must be explicit and non-empty.

TEXT:
"""

    SPELLING_RECALL_PROMPT = f"""
You are a strict Hindi spelling and orthography reviewer.

Scope:
- Review the full TEXT carefully from start to end
- Find clear spelling, orthography, matra, nukta, loanword-form, and punctuation-adjacent word errors that may have been missed in a first pass
- Do NOT translate, summarize, or rewrite
- Do NOT change meaning

Must-follow Hindi editorial rules:
- This publication style prefers non-chandrabindu normal forms such as "पांच" over "पाँच".
- This publication style also prefers the non-nukta house-style forms for ordinary Hindi words such as "ज्यादा", "जरूरी", "बाजार", and "नजर" unless a proper noun clearly requires nukta.
- For established loanwords that conventionally use the "ऑ" sound, prefer the standard spelling with "ऑ" only when the correction is genuinely unambiguous (for example, "कापी" -> "कॉपी", "कालेज" -> "कॉलेज"). Do not replace "काफी" with "कॉफी" unless the text clearly refers to the beverage.
- For transliterated proper nouns or foreign names ending in "िए", prefer the house-style "िये" ending when that is clearly the intended pronunciation/transliteration (for example, "तुर्किए" -> "तुर्किये"). Do not change ordinary Hindi verb forms such as "लिए", "दिए", "किए", "चाहिए", or "दीजिए".
- Do not create subjective wording changes.
- Original must be an exact substring from TEXT.

{rules_block}
{source_style_notes}

Return output strictly as table with header:
| Original | Corrected | Reason |

Reason must explicitly mention spelling/वर्तनी.

TEXT:
"""

    responses = []
    last_error = None

    for para in raw_paragraphs:
        try:
            out = generate_text(
                BASE_PROMPT + para,
                generation_config={
                    "temperature": 0,
                    "top_p": 1,
                    "top_k": 1,
                    "candidate_count": 1,
                    "max_output_tokens": 1400
                },
            )
            responses.append(out)
        except Exception as exc:
            last_error = exc
            continue

    for chunk in batch_hindi_texts(raw_paragraphs, max_chars=4200):
        try:
            out = generate_text(
                SPELLING_RECALL_PROMPT + chunk,
                generation_config={
                    "temperature": 0,
                    "top_p": 1,
                    "top_k": 1,
                    "candidate_count": 1,
                    "max_output_tokens": 1600
                },
            )
            responses.append(out)
        except Exception as exc:
            last_error = exc
            continue

    if not responses:
        return format_ai_error("grammar", last_error or RuntimeError("No Gemini response"))

    raw = "\n".join(responses)

    matches = re.findall(
        r"\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|",
        raw
    )

    rows = []
    seen = set()

    for o, c, r in matches:
        if o.lower() == "original" or c.lower() == "corrected" or r.lower() == "reason":
            continue
        if any(x.strip() in {"-", "--", "---"} for x in (o, c, r)):
            continue

        key = (canon_hi(o), canon_hi(c), canon_hi(r))
        if key in seen:
            continue

        seen.add(key)
        rows.append(f"| {o} | {c} | {r} |")

    if not rows:
        return ""

    return "\n".join(
        ["| Original | Corrected | Reason |",
         "|---|---|---|"] + rows
    )

# =================================================
# GEMINI EDITORIAL QC (HINDI GUIDELINES)
# =================================================
def gemini_editorial_review_hi(article_data, source_context=""):
    paragraphs = []
    for ctype, text in article_data:
        if ctype not in {"heading", "paragraph"}:
            continue
        paragraphs.extend(segment_hindi_review_text(text, max_chars=700))
    if not paragraphs:
        return ""
    source_style_notes = build_source_style_notes(source_context)

    base_prompt = """
You are an editorial QC reviewer for Hindi content.
Use English for Issue and Corrected Text. Keep fixes concise and specific.

Check for clear violations of these Hindi editorial rules:
- Use the Hindi danda "।" to end sentences (not a period).
- Do not add a danda to headlines, decks, labels, or subheadings that are not full sentences.
- Use double quotes for direct speech; single quotes for titles.
- Straight and curly variants of the same quote type are both acceptable; do not flag quote-shape-only swaps such as "..." vs “...”.
- This publication style avoids chandrabindu in normal house-style spellings; prefer forms like "पांच" over "पाँच".
- Do not suggest nukta-only rewrites in ordinary Hindi words when the non-nukta spelling is already acceptable house style (for example, ज्यादा/जरूरी/बाजार/नजर are acceptable without forcing nukta).
- For established loanwords with the "ऑ" sound, prefer standard spellings like "कॉपी" and "कॉलेज" only when the correction is genuinely unambiguous. Do not replace "काफी" with "कॉफी" unless the text clearly refers to the beverage.
- For transliterated proper nouns or foreign names ending in "िए", prefer the house-style "िये" ending when that is clearly the intended pronunciation/transliteration (for example, "तुर्किए" -> "तुर्किये"). Do not apply this to ordinary Hindi verb forms such as "लिए", "दिए", "किए", "चाहिए", or "दीजिए".
- Flag first-mention abbreviation issues only for named entities or terms that are genuinely unclear without expansion.
- Do not force expansion of common technical abbreviations, scientific labels, measurements, or UI labels.
- If a headline/subheading contains क्या/कैसे/क्यों/कब/कितना, it must end with "?".
- Use exactly three dots for ellipsis "...".
- Use numbers 1–9 in words, 10+ in numerals (except dates, time, prices, recipe ingredients).
- Do not use honorifics श्री/श्रीमती for names (only "महामहिम" for the President when needed).
- Do not enforce subjective style swaps, blanket comma preferences, or bans on sentence openings like "लेकिन"/"और".
- Do not translate acceptable English technical terms just to create a correction.
"""

    if source_style_notes:
        base_prompt += "\n" + source_style_notes.strip() + "\n"

    base_prompt += """

Rules for output:
- Excerpt must be an exact substring from the TEXT.
- Corrected Text must be the fixed version of Excerpt and must differ from Excerpt.
- If the text already follows the rule, do not flag it.
- For orthography / matra / chandrabindu / standard loanword-form fixes, describe the issue as spelling/वर्तनी, not grammar.
- Prefer existing non-nukta house-style spellings in ordinary Hindi words unless a proper noun or fixed foreign spelling clearly requires nukta.
- Do not stop after the first issue; identify all clear issues in the paragraph.

Return output strictly as a markdown table with header:
| Issue | Location | Excerpt | Corrected Text |

If there are no issues, return exactly one row:
| No issues found | - | - | - |
"""

    focused_prompt = """
You are a strict Hindi editorial QC reviewer. This is a focused recall pass.
Use English for Issue and Corrected Text.

Check only these categories, but identify all applicable issues from the paragraph:
- wrong quote type for direct speech (for example, single quotes instead of double quotes), but do not flag straight-vs-curly quote-shape differences
- sentence-ending punctuation or bracket spacing errors
- house-style orthography issues where chandrabindu should not be used
- established loanword spellings that clearly need the "ऑ" form
- abbreviation/acronym used before full form at first mention only when the abbreviation refers to a named entity and expansion is required for clarity
"""

    if source_style_notes:
        focused_prompt += "\n" + source_style_notes.strip() + "\n"

    focused_prompt += """

Rules for output:
- Excerpt must be an exact substring from the TEXT.
- Corrected Text must be the fixed version of Excerpt and must differ from Excerpt.
- Do not flag acceptable technical labels, measurement terms, or stylistic preferences.
- Return all clear issues, even if multiple rows come from the same paragraph.

Return output strictly as a markdown table with header:
| Issue | Location | Excerpt | Corrected Text |

If there are no issues, return exactly one row:
| No issues found | - | - | - |
"""

    responses = []
    last_error = None

    for i, para in enumerate(paragraphs, start=1):
        for prompt_template in (base_prompt, focused_prompt):
            prompt = (
                prompt_template
                + f"\nUse the literal location label `Paragraph {i}` for every row from this TEXT.\n\n"
                + "TEXT:\n"
                + para
            )
            try:
                responses.append(
                    generate_text(
                        prompt,
                        generation_config={
                            "temperature": 0,
                            "top_p": 1,
                            "top_k": 1,
                            "candidate_count": 1,
                            "max_output_tokens": 1400,
                        },
                    )
                )
            except Exception as exc:
                last_error = exc
                continue

    if not responses:
        return format_ai_error("editorial", last_error or RuntimeError("No Gemini response"))

    return "\n".join(responses)

# =================================================
# FACT STATEMENT HEURISTICS (EXPLICIT LAYER)
# =================================================
HINDI_FACT_VERBS = [
    "है", "था", "थे", "हैं",
    "कहा", "बताया", "बताई",
    "घोषणा", "जारी", "रिपोर्ट",
    "अनुसार", "मुताबिक"
]

def is_hindi_fact_sentence(sentence):
    return any(v in sentence for v in HINDI_FACT_VERBS)

def extract_fact_statements(article_data):
    statements = []
    seen = set()

    for ctype, text in article_data:
        if ctype not in {"paragraph", "table"}:
            continue

        for sent in split_hindi_sentences(text):
            if not is_hindi_fact_sentence(sent):
                continue

            key = canon_hi(sent)
            if key in seen:
                continue

            seen.add(key)
            statements.append(sent)

    return statements

# =================================================
# FACT CHECK (SECOND PASS, PARITY)
# =================================================
def article_hash(article_data):
    joined = "\n".join(t for _, t in article_data)
    return hashlib.md5(joined.encode("utf-8")).hexdigest()

def analysis_snapshot_key(article_data, source_context=""):
    return f"{PROMPT_VERSION_HI}:{get_domain(source_context)}:{article_hash(article_data)}"

def load_persistent_analysis_cache():
    try:
        with open(PERSISTENT_CACHE_PATH_HI, "r", encoding="utf-8") as f:
            data = json.load(f)
            return data if isinstance(data, dict) else {}
    except FileNotFoundError:
        return {}
    except Exception:
        return {}

def save_persistent_analysis_cache(cache):
    try:
        with open(PERSISTENT_CACHE_PATH_HI, "w", encoding="utf-8") as f:
            json.dump(cache, f, ensure_ascii=False)
    except Exception:
        pass

def load_analysis_snapshot(article_data, source_context=""):
    cache = load_persistent_analysis_cache()
    snapshot = cache.get(analysis_snapshot_key(article_data, source_context))
    return snapshot if snapshot_has_meaningful_output(snapshot) else None

def save_analysis_snapshot(article_data, snapshot, source_context=""):
    if not snapshot_has_meaningful_output(snapshot):
        return
    cache = load_persistent_analysis_cache()
    cache[analysis_snapshot_key(article_data, source_context)] = snapshot
    save_persistent_analysis_cache(cache)

def clear_persistent_analysis_cache():
    try:
        os.remove(PERSISTENT_CACHE_PATH_HI)
    except FileNotFoundError:
        pass
    except Exception:
        pass
    FACT_CACHE.clear()

def chunked(lst, size):
    for i in range(0, len(lst), size):
        yield lst[i:i + size]

def gemini_fact_check(article_data):
    key = article_hash(article_data)
    if key in FACT_CACHE:
        return FACT_CACHE[key]

    statements = extract_fact_statements(article_data)

    if not statements:
        return ""

    full_text = "\n".join(
        text for ctype, text in article_data if ctype in {"heading", "paragraph", "table"}
    )

    rows = []
    seen = set()
    had_success = False
    last_error = None

    for batch in chunked(statements, 5):
        block = "\n".join(f"- {s}" for s in batch)

        PROMPT = f"""
You are an internal factual consistency auditor.

Rules:
- Treat TEXT as closed
- No external knowledge
- Quote exact text
- No paraphrasing
- Only flag direct contradictions, impossible combinations, or statements that are unsupported by the article itself.
- Do not flag style, wording, naming preference, branding simplification, abbreviation expansion, punctuation, or grammar as factual issues.

Return table:
| Statement | Issue | Correct Fact |

TEXT:
{full_text}

STATEMENTS:
{block}
"""

        try:
            out = generate_text(
                PROMPT,
                generation_config={
                    "temperature": 0,
                    "top_p": 1,
                    "max_output_tokens": 512
                },
            )
            had_success = True
        except Exception as exc:
            last_error = exc
            continue

        matches = re.findall(
            r"\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|",
            out
        )

        for s, i, c in matches:
            if s.lower() == "statement" or i.lower() == "issue" or c.lower() == "correct fact":
                continue
            if any(x.strip() in {"-", "--", "---"} for x in (s, i, c)):
                continue
            if is_no_issue_fact(i, c):
                continue
            if is_style_only_fact(s, i, c):
                continue

            sig = (canon_hi(s), canon_hi(i))
            if sig in seen:
                continue

            seen.add(sig)
            rows.append(f"| {s} | {i} | {c} |")

        if len(seen) >= 10:
            break

    if not rows and not had_success:
        return format_ai_error("fact", last_error or RuntimeError("No Gemini response"))

    if not rows:
        return ""

    result = "\n".join(
        ["| Statement | Issue | Correct Fact |",
         "|---|---|---|"] + rows
    )

    FACT_CACHE[key] = result
    return result

@st.cache_data(show_spinner=False)
def cached_gemini_grammar_review(article_data, source_context=""):
    return gemini_grammar_review(article_data, source_context)

@st.cache_data(show_spinner=False)
def cached_gemini_editorial_review_hi(article_data, source_context=""):
    return gemini_editorial_review_hi(article_data, source_context)

@st.cache_data(show_spinner=False)
def cached_gemini_fact_check(article_data):
    return gemini_fact_check(article_data)

# =================================================
# PIPELINE
# =================================================

def render_ai_error(section_label: str, value: str):
    if is_ai_error_output(value):
        st.error(f"{section_label}: {value.replace('__ERROR__:', '').strip()}")
        return True
    return False

def run_pipeline(content):
    return content

# =================================================
# STREAMLIT UI (UNCHANGED STRUCTURE)
# =================================================
if not IMPORT_ONLY:
    st.sidebar.header("Input")
    source = st.sidebar.radio("Source", ["URL", "DOCX"])
    analyze_clicked = st.sidebar.button("Analyze")
    if st.sidebar.button("Clear cached AI outputs"):
        st.cache_data.clear()
        clear_persistent_analysis_cache()
        _clear_pending_analysis_state()
        for key in ("article_content", "input_key", "source_context", "source_label"):
            st.session_state.pop(key, None)

    article_content = None
    current_key = None
    source_context = ""

    if source == "URL":
        url = st.sidebar.text_input("Hindi Article URL")
        if url:
            current_key = f"url:{url.strip()}"
            source_context = url.strip()
        if analyze_clicked and url:
            article_content = clean_article(url)
            st.session_state["article_content"] = article_content
            st.session_state["input_key"] = current_key
            st.session_state["source_context"] = source_context
            st.session_state["source_label"] = url.strip()
            queue_analysis_run(
                "url",
                current_key,
                url.strip(),
            )
    else:
        uploaded = st.sidebar.file_uploader("Upload DOCX", type=["docx"])
        if uploaded:
            file_bytes = uploaded.getvalue()
            current_key = "docx:" + hashlib.sha256(file_bytes).hexdigest()
            if analyze_clicked:
                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as f:
                    f.write(file_bytes)
                    article_content = clean_docx(f.name)
                st.session_state["article_content"] = article_content
                st.session_state["input_key"] = current_key
                st.session_state["source_context"] = ""
                st.session_state["source_label"] = uploaded.name or current_key
                queue_analysis_run(
                    "docx",
                    current_key,
                    uploaded.name or current_key,
                )

    if article_content is None:
        if current_key and st.session_state.get("input_key") == current_key:
            article_content = st.session_state.get("article_content")
            source_context = st.session_state.get("source_context", source_context)

    source_label = st.session_state.get("source_label", "")

    if article_content:
        qc_content = run_pipeline(article_content)

        st.subheader("📄 Final Article")
        for _, t in qc_content:
            st.write(t)

        st.divider()

        st.subheader("🤖 Gemini QC Review")

        report_pdf_bytes = None
        report_pdf_error = None

        article_text = "\n".join(
            t for c, t in article_content if c in {"heading", "paragraph", "table"}
        )

        snapshot = load_analysis_snapshot(qc_content, source_context)
        if snapshot:
            raw = snapshot.get("grammar_raw", "")
            editorial_raw = snapshot.get("editorial_raw", "")
            fact_result = snapshot.get("fact_result", "")
        else:
            raw = cached_gemini_grammar_review(qc_content, source_context)
            editorial_raw = cached_gemini_editorial_review_hi(qc_content, source_context)
            fact_result = cached_gemini_fact_check(qc_content)
            save_analysis_snapshot(
                qc_content,
                {
                    "grammar_raw": raw,
                    "editorial_raw": editorial_raw,
                    "fact_result": fact_result,
                },
                source_context,
            )

        clean = filter_gemini_rows(raw, article_text)
        language_rows = parse_language_rows(clean, qc_content)

        editorial_clean = filter_editorial_rows(editorial_raw, article_text)
        editorial_rows = parse_editorial_rows(editorial_clean, qc_content)
        editorial_language_rows = parse_editorial_as_language_rows(editorial_clean, qc_content)
        rule_based_rows = rule_based_spelling_rows(qc_content)
        editorial_display = build_editorial_table(editorial_rows)

        spelling_table, grammar_table = build_language_tables(
            language_rows + rule_based_rows,
            editorial_language_rows,
        )
        spelling_count = len(spelling_table)
        grammar_count = len(grammar_table)
        editorial_count = len(editorial_rows)
        fact_count = 0 if is_ai_error_output(fact_result) else count_markdown_rows(fact_result, "Statement")
        render_qc_score_summary(
            spelling_count,
            grammar_count,
            editorial_count,
            fact_count,
            any(is_ai_error_output(value) for value in (raw, editorial_raw, fact_result)),
        )

        st.markdown("### ✍️ Spelling Issues")
        if render_ai_error("Spelling/Grammar AI", raw):
            pass
        elif spelling_table:
            st.markdown(render_language_table(spelling_table), unsafe_allow_html=True)
        else:
            st.success("✅ No spelling issues found")

        st.markdown("### 🧠 Grammar Issues")
        if render_ai_error("Spelling/Grammar AI", raw):
            pass
        elif grammar_table:
            st.markdown(render_language_table(grammar_table), unsafe_allow_html=True)
        else:
            st.success("✅ No grammar issues found")

        st.markdown("### 🧠 Gemini Editorial Review")
        if render_ai_error("Editorial AI", editorial_raw):
            pass
        elif editorial_display:
            st.markdown(editorial_display)
        else:
            st.success("✅ No editorial issues found")

        st.markdown("### 📌 Fact Check")
        if render_ai_error("Fact-check AI", fact_result):
            pass
        elif not fact_result or "| Statement |" not in fact_result:
            st.success("✅ No factual issues found")
        else:
            st.markdown(fact_result)

        report_pdf_bytes, report_pdf_error = build_hindi_qc_report_pdf(
            source_label or (url.strip() if source == "URL" and url else current_key or "QC Report"),
            _current_access_email(),
            spelling_table,
            grammar_table,
            editorial_rows,
            fact_result,
        )

        if report_pdf_bytes:
            st.download_button(
                "Download QC Report (PDF)",
                data=report_pdf_bytes,
                file_name=f"hindi_qc_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                mime="application/pdf",
            )
        elif report_pdf_error:
            st.caption(f"PDF report unavailable: {report_pdf_error}")

        log_analysis_run(
            "hindi_qc",
            _current_access_email(),
            st.session_state.get("_pending_source_type", source.lower()),
            st.session_state.get("_pending_source_identity", current_key or ""),
            st.session_state.get("_pending_source_label", url.strip() if source == "URL" and url else ""),
            st.session_state.get("_pending_analysis_key", ""),
            spelling_count,
            grammar_count,
            editorial_count,
            fact_count,
        )

    if _is_admin_user():
        render_admin_dashboard("hindi_qc")
